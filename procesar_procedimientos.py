
import os
from docx import Document
import pandas as pd
import re
import gc
from zipfile import BadZipFile

def extraer_datos_encabezado(document):
    try:
        header = document.sections[0].header
        tables = header.tables
        if not tables:
            return None

        tabla = tables[0]
        codigo = tabla.cell(0, 2).text.upper().strip().replace("CÓDIGO:", "").replace("CODIGO:", "").strip()
        version = tabla.cell(1, 2).text.upper().strip().replace("VERSIÓN:", "").replace("VERSION:", "").strip()
        nombre = tabla.cell(2, 1).text.strip()
        edicion = tabla.cell(2, 2).text.upper().strip().replace("EDICIÓN:", "").replace("EDICION:", "").strip()

        return {
            "Código": codigo,
            "Nombre del procedimiento": nombre,
            "Versión": version,
            "Edición": edicion
        }
    except Exception as e:
        print(f"Error procesando encabezado: {e}")
        return {}

def detectar_secciones_principales(document):
    indices = {}
    secciones_buscar = [
        "INFORMACIÓN GENERAL DEL PROCEDIMIENTO",
        "PELIGROS, RIESGOS Y CONTROLES DE LA ACTIVIDAD",
        "ASPECTOS E IMPACTOS AMBIENTALES Y CONTROLES DE LA ACTIVIDAD",
        "CONDICIONES PREVIAS A LA EJECUCION DE LA ACTIVIDAD",
        "DESCRIPCIÓN DE ACTIVIDADES",
        "CONSIDERACIONES POSTERIORES A LA EJECUCIÓN DE LA ACTIVIDAD"
    ]

    for i, para in enumerate(document.paragraphs):
        texto_normalizado = ' '.join(para.text.strip().split()).upper()
        for seccion in secciones_buscar:
            seccion_normalizada = ' '.join(seccion.strip().split()).upper()
            if seccion_normalizada in texto_normalizado:
                indices[seccion] = i
                break

    return indices

def extraer_seccion_info_general(document, indice_inicio, indice_fin):
    patrones_subseccion = [
        (r"OBJETO", "OBJETO"),
        (r"ALCANCE", "ALCANCE"),
        (r"DISCIPLINA", "DISCIPLINA"),
        (r"RECURSOS REQUERIDOS", "RECURSOS_REQUERIDOS"),
        (r"ELEMENTOS PROTECCION PERSONAL", "ELEMENTOS_PROTECCION")
    ]

    info_general = {}
    subseccion_actual = None

    for i in range(indice_inicio + 1, indice_fin):
        texto = document.paragraphs[i].text.strip()
        if not texto:
            continue

        es_subseccion = False
        for patron, clave in patrones_subseccion:
            if re.search(patron, texto):
                subseccion_actual = clave
                info_general[subseccion_actual] = ""
                es_subseccion = True
                break

        if not es_subseccion and subseccion_actual:
            if info_general[subseccion_actual]:
                info_general[subseccion_actual] += "\n" + texto
            else:
                info_general[subseccion_actual] = texto

    return info_general

def extraer_texto_completo_seccion(document, indice_inicio, indice_fin):
    texto_completo = []
    for i in range(indice_inicio + 1, indice_fin):
        texto = document.paragraphs[i].text.strip()
        if texto:
            texto_completo.append(texto)
    return "\n".join(texto_completo)

def procesar_documento(ruta_archivo):
    try:
        doc = Document(ruta_archivo)
        datos = extraer_datos_encabezado(doc) or {}
        datos["Nombre del archivo"] = os.path.splitext(os.path.basename(ruta_archivo))[0]

        indices = detectar_secciones_principales(doc)

        if "INFORMACIÓN GENERAL DEL PROCEDIMIENTO" in indices and "PELIGROS, RIESGOS Y CONTROLES DE LA ACTIVIDAD" in indices:
            info_general = extraer_seccion_info_general(
                doc,
                indices["INFORMACIÓN GENERAL DEL PROCEDIMIENTO"],
                indices["PELIGROS, RIESGOS Y CONTROLES DE LA ACTIVIDAD"]
            )
            for subseccion, contenido in info_general.items():
                datos[f"InfoGeneral_{subseccion}"] = contenido

        if "DESCRIPCIÓN DE ACTIVIDADES" in indices and "CONSIDERACIONES POSTERIORES A LA EJECUCIÓN DE LA ACTIVIDAD" in indices:
            descripcion_texto = extraer_texto_completo_seccion(
                doc,
                indices["DESCRIPCIÓN DE ACTIVIDADES"],
                indices["CONSIDERACIONES POSTERIORES A LA EJECUCIÓN DE LA ACTIVIDAD"]
            )
            datos["Descripcion_Actividades"] = descripcion_texto

        return datos

    except BadZipFile:
        print(f"❌ Archivo corrupto o dañado: {ruta_archivo}")
        return {"Nombre del archivo": os.path.splitext(os.path.basename(ruta_archivo))[0], "Error": "Archivo corrupto o dañado"}

    except Exception as e:
        print(f"Error procesando documento {ruta_archivo}: {e}")
        return {"Nombre del archivo": os.path.splitext(os.path.basename(ruta_archivo))[0], "Error": str(e)}

def procesar_documentos_en_carpeta(carpeta, master_file):
    datos = []
    for archivo in os.listdir(carpeta):
        if archivo.endswith(".docx"):
            ruta = os.path.join(carpeta, archivo)
            try:
                datos_doc = procesar_documento(ruta)
                datos.append(datos_doc)
                print(f"Procesado: {archivo}")
            except Exception as e:
                print(f"Error leyendo {archivo}: {e}")

    df = pd.DataFrame(datos)
    df["check"] = df.apply(lambda row: 1 if row.get("Nombre del archivo") == row.get("Código") else 0, axis=1)
    columnas = ["Nombre del archivo", "Código", "check", "Nombre del procedimiento"] +                [col for col in df.columns if col not in ["Nombre del archivo", "Código", "check", "Nombre del procedimiento"]]
    df = df[columnas]

    # Load master file and merge
    master_df = pd.read_excel(master_file, engine='openpyxl')
    df = df.merge(master_df[["Codigo", "Plantilla", "Tipo de Procedimiento", "Campo"]],
                  left_on="Código", right_on="Codigo", how="left")
    df.drop(columns=["Codigo"], inplace=True)

    # Check if Matriz Relacional.xlsx exists
    if os.path.exists("Matriz Relacional.xlsx"):
        existing_df = pd.read_excel("Matriz Relacional.xlsx", engine='openpyxl')
        combined_df = pd.concat([existing_df, df], ignore_index=True)
        combined_df.drop_duplicates(subset=["Nombre del archivo", "Versión", "Edición"], keep='last', inplace=True)
        combined_df.to_excel("Matriz Relacional.xlsx", index=False)
        print("✅ Actualización completada. Archivo 'Matriz Relacional.xlsx' actualizado.")
    else:
        df.to_excel("Matriz Relacional.xlsx", index=False)
        print("✅ Extracción completada. Archivo 'Matriz Relacional.xlsx' generado.")

    # Liberar memoria
    del datos, df, master_df, combined_df
    gc.collect()
