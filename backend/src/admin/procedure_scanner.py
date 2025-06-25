"""
Scanner de procedimientos para el módulo administrativo de InemecTest
Basado en la lógica de word_to_excel.py con integración al sistema existente
"""

import os
import json
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
import hashlib
from docx import Document
import pandas as pd
from .utils import create_tracking_key
from .utils import extract_procedure_code_and_version
from .config import get_admin_file_path, get_admin_directory_path

from ..config import (
    get_data_file_path,
    DATA_SHEETS,
    PROCEDURES_COLUMNS,
    ensure_data_directory
)

class ProcedureScanner:
    """
    Scanner de procedimientos que detecta archivos .docx nuevos o actualizados
    y determina cuáles necesitan generación de preguntas
    """
    
    def __init__(self, procedures_source_dir: str | None = None, max_workers: int = 4):
        """
        Inicializar scanner
        
        Args:
            procedures_source_dir: Directorio donde están los archivos .docx de procedimientos
            max_workers: Número máximo de hilos para procesamiento paralelo
        """

        if procedures_source_dir is None:
            procedures_source_dir = str(get_admin_directory_path("procedures_source"))
        self.procedures_source_dir = Path(procedures_source_dir)
        self.excel_file = get_data_file_path()
        self.tracking_file = get_admin_file_path("tracking")
        self.cache_file = get_admin_file_path("metadata_cache")
        self.max_workers = max_workers
        
        # Crear directorios si no existen
        ensure_data_directory()
        self.procedures_source_dir.mkdir(exist_ok=True)
        
        print(f"📁 Scanner inicializado:")
        print(f"   - Directorio fuente: {self.procedures_source_dir}")
        print(f"   - Archivo Excel: {self.excel_file}")
        print(f"   - Tracking file: {self.tracking_file}")
        print(f"   - Cache file: {self.cache_file}")
        print(f"   - Max workers: {self.max_workers}")
    
    def extraer_datos_encabezado(self, document: Document) -> Dict[str, Any]:
        """
        Extrae información del encabezado del documento
        Basado en word_to_excel.py
        """
        try:
            header = document.sections[0].header
            tables = header.tables
            if not tables:
                return {}
            
            tabla = tables[0]
            codigo = tabla.cell(0, 2).text.strip().replace("CÓDIGO:", "").strip()
            version = tabla.cell(1, 2).text.strip().replace("VERSIÓN:", "").strip()
            nombre = tabla.cell(2, 1).text.strip()
            edicion = tabla.cell(2, 2).text.strip().replace("EDICIÓN:", "").strip()
            
            return {
                "codigo": codigo,
                "nombre": nombre,
                "version": version,
                "edicion": edicion
            }
        except Exception as e:
            print(f"⚠️ Error procesando encabezado: {e}")
            return {}
    
    def detectar_secciones_principales(self, document: Document) -> Dict[str, int]:
        """
        Detecta los índices de inicio y fin de las secciones principales
        Basado en word_to_excel.py
        """
        indices = {}
        secciones_buscar = [
            "INFORMACIÓN GENERAL DEL PROCEDIMIENTO",
            "PELIGROS, RIESGOS Y CONTROLES DE LA ACTIVIDAD",
            "ASPECTOS E IMPACTOS AMBIENTALES Y CONTROLES DE LA ACTIVIDAD",
            "CONDICIONES PREVIAS A LA EJECUCION DE LA ACTIVIDAD",
            "DESCRIPCIÓN DE ACTIVIDADES",
            "CONSIDERACIONES POSTERIORES A LA EJECUCIÓN DE LA ACTIVIDAD"
        ]
        
        for i, para in enumerate(document.paragraphs):
            # Normalizar el texto para la comparación
            texto_normalizado = ' '.join(para.text.strip().split()).upper()
            
            for seccion in secciones_buscar:
                seccion_normalizada = ' '.join(seccion.strip().split()).upper()
                if seccion_normalizada in texto_normalizado:
                    indices[seccion] = i
                    break
        
        return indices
    
    def extraer_seccion_info_general(self, document: Document, indice_inicio: int, indice_fin: int) -> Dict[str, str]:
        """
        Extrae la información de la sección INFORMACIÓN GENERAL
        Basado en word_to_excel.py
        """
        patrones_subseccion = [
            (r"OBJETO", "OBJETO"),
            (r"ALCANCE", "ALCANCE"),
            (r"DISCIPLINA", "DISCIPLINA"),
            (r"RECURSOS REQUERIDOS", "RECURSOS_REQUERIDOS"),
            (r"ELEMENTOS PROTECCION PERSONAL", "ELEMENTOS_PROTECCION")
        ]
        
        info_general = {}
        subseccion_actual = None
        
        for i in range(indice_inicio + 1, indice_fin):
            if i >= len(document.paragraphs):
                break
                
            texto = document.paragraphs[i].text.strip()
            if not texto:
                continue
            
            # Verificar si es un título de subsección
            es_subseccion = False
            for patron, clave in patrones_subseccion:
                if re.search(patron, texto, re.IGNORECASE):
                    subseccion_actual = clave
                    info_general[subseccion_actual] = ""
                    es_subseccion = True
                    break
            
            # Si no es subsección y tenemos subsección actual, añadir contenido
            if not es_subseccion and subseccion_actual:
                if info_general[subseccion_actual]:
                    info_general[subseccion_actual] += "\n" + texto
                else:
                    info_general[subseccion_actual] = texto
        
        return info_general
    
    def extraer_texto_completo_seccion(self, document: Document, indice_inicio: int, indice_fin: int) -> str:
        """
        Extrae todo el texto entre dos índices de párrafo como un solo bloque
        Basado en word_to_excel.py
        """
        texto_completo = []
        
        for i in range(indice_inicio + 1, indice_fin):
            if i >= len(document.paragraphs):
                break
                
            texto = document.paragraphs[i].text.strip()
            if texto:
                texto_completo.append(texto)
        
        return "\n".join(texto_completo)
    
    
    def procesar_documento(self, ruta_archivo: Path) -> Dict[str, Any]:
        """
        Procesa un documento .docx y extrae toda la información relevante
        🔧 VERSIÓN MEJORADA: Integra lógica de procesar_procedimientos.py
        """
        try:
            print(f"📄 [DEBUG] Intentando abrir documento: {ruta_archivo}")
            doc = Document(ruta_archivo)
            print(f"📄 [DEBUG] Documento abierto exitosamente")
            
            # ✅ USAR SOLO FILENAME - Más confiable
            codigo_raw, version_final = extract_procedure_code_and_version(ruta_archivo.name)
            
            # ✅ LIMPIAR el código de prefijos no deseados
            codigo_final = self._limpiar_codigo(codigo_raw)
            
            # Extraer datos del encabezado con lógica mejorada
            datos_encabezado = self.extraer_datos_encabezado(doc)
            nombre_encabezado = datos_encabezado.get("nombre", "")
            edicion_encabezado = datos_encabezado.get("edicion", "")
            
            # 🎯 ESTRUCTURA COMPLETA usando procesar_procedimientos.py
            datos = {
                "codigo": codigo_final,
                "version": str(version_final),
                "nombre": nombre_encabezado or f"Procedimiento {codigo_final}",
                "alcance": "",
                "objetivo": "",
                "archivo": ruta_archivo.name,
                "ruta_completa": str(ruta_archivo),
                "fecha_escaneado": datetime.now().isoformat(),
                "edicion": edicion_encabezado,
                "disciplina": "",
                "recursos_requeridos": "",
                "elementos_proteccion": "",
                "descripcion_actividades": "",
                "tipo_procedimiento": "",
                "campo": ""
            }
            
            # 🔍 DEBUG: Mostrar datos procesados
            print(f"📄 Procesado: {ruta_archivo.name}")
            print(f"   - Código original: {codigo_raw}")
            print(f"   - Código limpio: {codigo_final}")
            print(f"   - Versión: {version_final}")
            print(f"   - Tracking key sería: {codigo_final}_v{version_final}")
            
            # Extraer secciones con lógica mejorada
            indices = self.detectar_secciones_principales(doc)
            
            # Información General del Procedimiento
            if "INFORMACIÓN GENERAL DEL PROCEDIMIENTO" in indices and "PELIGROS, RIESGOS Y CONTROLES DE LA ACTIVIDAD" in indices:
                info_general = self.extraer_seccion_info_general(
                    doc, 
                    indices["INFORMACIÓN GENERAL DEL PROCEDIMIENTO"],
                    indices["PELIGROS, RIESGOS Y CONTROLES DE LA ACTIVIDAD"]
                )
                datos["alcance"] = info_general.get("ALCANCE", "")
                datos["objetivo"] = info_general.get("OBJETO", "")
                datos["disciplina"] = info_general.get("DISCIPLINA", "")
                datos["recursos_requeridos"] = info_general.get("RECURSOS_REQUERIDOS", "")
                datos["elementos_proteccion"] = info_general.get("ELEMENTOS_PROTECCION", "")
            
            # Descripción de Actividades
            if "DESCRIPCIÓN DE ACTIVIDADES" in indices and "CONSIDERACIONES POSTERIORES A LA EJECUCIÓN DE LA ACTIVIDAD" in indices:
                descripcion_texto = self.extraer_texto_completo_seccion(
                    doc,
                    indices["DESCRIPCIÓN DE ACTIVIDADES"],
                    indices["CONSIDERACIONES POSTERIORES A LA EJECUCIÓN DE LA ACTIVIDAD"]
                )
                datos["descripcion_actividades"] = descripcion_texto
            
            # Detectar tipo de procedimiento basado en contenido
            datos["tipo_procedimiento"] = self._detectar_tipo_procedimiento(doc, datos)
            
            # Campo se puede derivar del código o contenido
            datos["campo"] = self._detectar_campo_procedimiento(codigo_final, datos)
            
            return datos
            
        except Exception as e:
            print(f"❌ Error procesando documento {ruta_archivo}: {e}")
            codigo_fallback, version_fallback = extract_procedure_code_and_version(ruta_archivo.name)
            codigo_fallback = self._limpiar_codigo(codigo_fallback)
            return {
                "codigo": codigo_fallback,
                "version": str(version_fallback),
                "nombre": f"ERROR: {ruta_archivo.name}",
                "alcance": "",
                "objetivo": "",
                "archivo": ruta_archivo.name,
                "ruta_completa": str(ruta_archivo),
                "error": str(e),
                "fecha_escaneado": datetime.now().isoformat(),
                "edicion": "",
                "disciplina": "",
                "recursos_requeridos": "",
                "elementos_proteccion": "",
                "descripcion_actividades": "",
                "tipo_procedimiento": "",
                "campo": ""
            }

    def _limpiar_codigo(self, codigo: str) -> str:
        """
        Limpiar código de prefijos no deseados
        """
        import re
        
        # Remover prefijos comunes
        prefijos_a_remover = ["CODIGO:", "CÓDIGO:", "CODIGO ", "CÓDIGO "]
        
        codigo_limpio = codigo.strip()
        
        for prefijo in prefijos_a_remover:
            if codigo_limpio.upper().startswith(prefijo.upper()):
                codigo_limpio = codigo_limpio[len(prefijo):].strip()
                break
        
        # Validar formato final
        if not re.match(r'^PEP-PRO-\d+$', codigo_limpio):
            print(f"⚠️ Código limpiado no válido: '{codigo_limpio}' (original: '{codigo}')")
            # Intentar extraer PEP-PRO-XXX del texto
            pep_match = re.search(r'(PEP-PRO-\d+)', codigo_limpio)
            if pep_match:
                codigo_limpio = pep_match.group(1)
                print(f"   ✅ Código extraído: {codigo_limpio}")
        
        return codigo_limpio
    
    def _detectar_tipo_procedimiento(self, document: Document, datos: Dict[str, Any]) -> str:
        """
        Detectar tipo de procedimiento basado en contenido del documento
        """
        # Buscar en el encabezado y contenido
        texto_completo = ""
        
        # Buscar en encabezado
        try:
            header = document.sections[0].header
            if header.tables:
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            texto_completo += cell.text.upper() + " "
        except:
            pass
        
        # Buscar en primeros párrafos
        for i, para in enumerate(document.paragraphs[:20]):  # Solo primeros 20 párrafos
            texto_completo += para.text.upper() + " "
        
        # Patrones para detectar tipo
        if any(keyword in texto_completo for keyword in ["OPERATIVO", "OPERACIÓN", "OPERADOR"]):
            return "OPERATIVO"
        elif any(keyword in texto_completo for keyword in ["TÉCNICO", "TECNICO", "MANTENIMIENTO", "REPARACIÓN"]):
            return "TECNICO"
        elif any(keyword in texto_completo for keyword in ["ADMINISTRATIVO", "GESTIÓN", "ADMINISTRACIÓN"]):
            return "ADMINISTRATIVO"
        
        # Default basado en código
        if "PRO" in datos["codigo"]:
            return "TECNICO"
        
        return "TECNICO"  # Default
    
    def _detectar_campo_procedimiento(self, codigo: str, datos: Dict[str, Any]) -> str:
        """
        Detectar campo del procedimiento basado en código y contenido
        """
        # Mapeo común de códigos a campos
        codigo_upper = codigo.upper()
        
        # Campos conocidos basados en patrones comunes en la industria
        if any(keyword in codigo_upper for keyword in ["CUPIAGUA", "CUP"]):
            return "Campo Cupiagua"
        elif any(keyword in codigo_upper for keyword in ["CUSIANA", "CUS"]):
            return "Campo Cusiana"
        elif any(keyword in codigo_upper for keyword in ["FLOREÑA", "FLO"]):
            return "Campo Floreña"
        elif any(keyword in codigo_upper for keyword in ["PAUTO", "PAU"]):
            return "Campo Pauto"
        
        # Buscar en contenido
        contenido_buscar = (datos.get("nombre", "") + " " + 
                           datos.get("alcance", "") + " " + 
                           datos.get("objetivo", "")).upper()
        
        if "CUPIAGUA" in contenido_buscar:
            return "Campo Cupiagua"
        elif "CUSIANA" in contenido_buscar:
            return "Campo Cusiana"
        elif "FLOREÑA" in contenido_buscar:
            return "Campo Floreña"
        elif "PAUTO" in contenido_buscar:
            return "Campo Pauto"
        
        # Default genérico
        return "Campo General"
    
    def cargar_tracking_data(self) -> Dict[str, Any]:
        """
        Cargar datos de tracking de generación de preguntas
        """
        # Estructura por defecto
        default_structure = {
            "generated_questions": {},  # {codigo_version: {...}}
            "last_scan": None,
            "scan_history": []
        }
        
        if self.tracking_file.exists():
            try:
                with open(self.tracking_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Asegurar que tiene todas las claves necesarias
                for key, default_value in default_structure.items():
                    if key not in data:
                        data[key] = default_value
                        print(f"⚠️ [DEBUG] Agregando clave faltante '{key}' al tracking data")
                
                return data
            except Exception as e:
                print(f"⚠️ Error cargando tracking data: {e}")
        
        return default_structure
    
    def guardar_tracking_data(self, tracking_data: Dict[str, Any]):
        """
        Guardar datos de tracking
        """
        try:
            with open(self.tracking_file, 'w', encoding='utf-8') as f:
                json.dump(tracking_data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"❌ Error guardando tracking data: {e}")
    
    
    def escanear_directorio(self) -> Dict[str, Any]:
        """
        Escanear directorio de procedimientos y detectar cambios (VERSIÓN OPTIMIZADA)
        """
        print(f"🚀 [OPTIMIZED] Escaneando directorio: {self.procedures_source_dir}")
        print(f"🚀 [OPTIMIZED] Usando {self.max_workers} hilos para procesamiento paralelo")
        
        # Cargar tracking data y cache
        try:
            tracking_data = self.cargar_tracking_data()
            cache = self.cargar_metadata_cache()
            print(f"🔍 [DEBUG] Tracking data: {len(tracking_data.get('generated_questions', {}))} items")
            print(f"💾 [DEBUG] Cache entries: {len(cache)} items")
        except Exception as e:
            print(f"❌ [DEBUG] Error cargando datos: {e}")
            tracking_data = {"generated_questions": {}, "last_scan": None, "scan_history": []}
            cache = {}
        
        if not self.procedures_source_dir.exists():
            print(f"❌ [DEBUG] Directorio no existe: {self.procedures_source_dir}")
            return {
                "success": False,
                "message": f"Directorio no existe: {self.procedures_source_dir}",
                "archivos_encontrados": 0,
                "procedimientos_nuevos": 0,
                "procedimientos_actualizados": 0,
                "cola_generacion": []
            }
        
        try:
            # Encontrar archivos .docx válidos
            print(f"🔍 [DEBUG] Buscando archivos *.docx...")
            docx_files = [
                archivo for archivo in self.procedures_source_dir.glob("*.docx")
                if not archivo.name.startswith("~$")  # Filtrar archivos temporales
            ]
            print(f"📄 [DEBUG] Archivos válidos encontrados: {len(docx_files)}")
            
            if not docx_files:
                print(f"⚠️ No se encontraron archivos .docx válidos")
                return {
                    "success": True,
                    "message": "No se encontraron archivos .docx",
                    "archivos_encontrados": 0,
                    "procedimientos_nuevos": 0,
                    "procedimientos_actualizados": 0,
                    "cola_generacion": []
                }
            
            # Procesar archivos en lotes paralelos
            batch_size = max(10, len(docx_files) // self.max_workers)
            batches = [docx_files[i:i + batch_size] for i in range(0, len(docx_files), batch_size)]
            
            print(f"🔄 [PROCESSING] Procesando {len(docx_files)} archivos en {len(batches)} lotes")
            
            procedimientos_escaneados = []
            
            for i, batch in enumerate(batches, 1):
                print(f"📦 [BATCH {i}/{len(batches)}] Procesando {len(batch)} archivos...")
                resultados_batch = self.procesar_lote_archivos(batch, cache)
                procedimientos_escaneados.extend(resultados_batch)
                print(f"✅ [BATCH {i}/{len(batches)}] Completados {len(resultados_batch)} archivos")
            
            # Guardar cache actualizado
            self.guardar_metadata_cache(cache)
            print(f"💾 Cache actualizado con {len(cache)} entradas")
        
        except Exception as scan_error:
            print(f"❌ [DEBUG] Error escaneando directorio: {scan_error}")
            return {
                "success": False,
                "message": f"Error escaneando directorio: {scan_error}",
                "archivos_encontrados": 0,
                "procedimientos_nuevos": 0,
                "procedimientos_actualizados": 0,
                "cola_generacion": []
            }
            
        archivos_encontrados = [proc["archivo"] for proc in procedimientos_escaneados]
        
        print(f"📄 [DEBUG] Archivos .docx válidos encontrados: {len(archivos_encontrados)}")
        print(f"📄 [DEBUG] Procedimientos procesados: {len(procedimientos_escaneados)}")
        
        # Analizar todos los procedimientos y determinar su estado
        cola_generacion = []
        procedimientos_nuevos = 0
        procedimientos_ya_procesados = 0
        
        for proc_data in procedimientos_escaneados:
            codigo = proc_data["codigo"]
            version = proc_data.get("version", "1")
            
            tracking_key = create_tracking_key(codigo, version)

            # Determinar estado del procedimiento
            estado = "nuevo"  # Por defecto es nuevo
            generated_questions_count = 0
            
            if tracking_key in tracking_data.get("generated_questions", {}):
                tracking_info = tracking_data["generated_questions"][tracking_key]
                status = tracking_info.get("status")
                if status == "completed":
                    estado = "ya_procesado"
                    generated_questions_count = tracking_info.get("preguntas_count", 0)
                    procedimientos_ya_procesados += 1
                else:
                    # Estado intermedio (failed, processing, etc.)
                    estado = "necesita_reproceso"
                    procedimientos_nuevos += 1
            else:
                # No existe en tracking, es nuevo
                procedimientos_nuevos += 1
            
            print(f"🔍 [DEBUG] Procedimiento {codigo} v{version}: estado={estado}, preguntas={generated_questions_count}")
            
            # Añadir TODOS los procedimientos a la cola (con su estado correspondiente)
            item_cola = {
                "codigo": codigo,
                "nombre": proc_data["nombre"],
                "version": version,
                "archivo": proc_data["archivo"],
                "estado": estado,
                "tracking_key": tracking_key,
                "datos_completos": proc_data,
                "preguntas_generadas": generated_questions_count,
                "puede_generar": estado in ["nuevo", "necesita_reproceso"],
                "puede_regenerar": estado == "ya_procesado"
            }
            
            cola_generacion.append(item_cola)
        
        # Actualizar tracking data
        scan_info = {
            "timestamp": datetime.now().isoformat(),
            "archivos_escaneados": len(archivos_encontrados),
            "procedimientos_en_cola": len(cola_generacion)
        }
        
        tracking_data["last_scan"] = scan_info
        
        # Asegurar que scan_history existe
        if "scan_history" not in tracking_data:
            tracking_data["scan_history"] = []
        
        tracking_data["scan_history"].append(scan_info)
        
        # Mantener solo últimos 10 scans en historial
        if len(tracking_data["scan_history"]) > 10:
            tracking_data["scan_history"] = tracking_data["scan_history"][-10:]
        
        self.guardar_tracking_data(tracking_data)
        
        resultado = {
            "success": True,
            "message": f"Escaneo completado: {len(archivos_encontrados)} archivos encontrados",
            "archivos_encontrados": len(archivos_encontrados),
            "procedimientos_nuevos": procedimientos_nuevos,
            "procedimientos_ya_procesados": procedimientos_ya_procesados,
            "total_procedimientos": len(cola_generacion),
            "cola_generacion": cola_generacion,
            "tracking_file": str(self.tracking_file),
            "timestamp": datetime.now().isoformat()
        }
        
        print(f"✅ Escaneo completado:")
        print(f"   - Archivos encontrados: {len(archivos_encontrados)}")
        print(f"   - Procedimientos nuevos: {procedimientos_nuevos}")
        print(f"   - Procedimientos ya procesados: {procedimientos_ya_procesados}")
        print(f"   - Total en cola: {len(cola_generacion)}")
        
        return resultado
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Obtener hash del archivo basado en su ruta y fecha de modificación"""
        try:
            stat = file_path.stat()
            content = f"{file_path.name}_{stat.st_mtime}_{stat.st_size}"
            return hashlib.md5(content.encode()).hexdigest()
        except:
            return hashlib.md5(str(file_path).encode()).hexdigest()
    
    def cargar_metadata_cache(self) -> Dict[str, Dict[str, Any]]:
        """Cargar cache de metadatos de archivos"""
        if self.cache_file.exists():
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"⚠️ Error cargando cache: {e}")
        return {}
    
    def guardar_metadata_cache(self, cache: Dict[str, Dict[str, Any]]):
        """Guardar cache de metadatos"""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(cache, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"❌ Error guardando cache: {e}")
    
    def procesar_documento_con_cache(self, ruta_archivo: Path, cache: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
        """Procesar documento usando cache si está disponible"""
        file_hash = self._get_file_hash(ruta_archivo)
        
        # Verificar si está en cache y es válido
        if file_hash in cache:
            cached_data = cache[file_hash]
            if cached_data.get('archivo') == ruta_archivo.name:
                print(f"💾 [CACHE HIT] {ruta_archivo.name}")
                return cached_data
        
        # Procesar documento
        print(f"🔄 [PROCESSING] {ruta_archivo.name}")
        datos = self.procesar_documento(ruta_archivo)
        
        # Guardar en cache
        cache[file_hash] = datos
        
        return datos
    
    def procesar_lote_archivos(self, archivos: List[Path], cache: Dict[str, Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Procesar un lote de archivos en paralelo"""
        resultados = []
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Enviar tareas
            future_to_archivo = {
                executor.submit(self.procesar_documento_con_cache, archivo, cache): archivo 
                for archivo in archivos
            }
            
            # Recoger resultados
            for future in as_completed(future_to_archivo):
                archivo = future_to_archivo[future]
                try:
                    resultado = future.result()
                    resultados.append(resultado)
                    print(f"✅ Completado: {archivo.name}")
                except Exception as e:
                    print(f"❌ Error procesando {archivo.name}: {e}")
                    # Crear resultado de error
                    codigo_fallback, version_fallback = extract_procedure_code_and_version(archivo.name)
                    codigo_fallback = self._limpiar_codigo(codigo_fallback)
                    resultado_error = {
                        "codigo": codigo_fallback,
                        "version": str(version_fallback),
                        "nombre": f"ERROR: {archivo.name}",
                        "alcance": "",
                        "objetivo": "",
                        "archivo": archivo.name,
                        "ruta_completa": str(archivo),
                        "error": str(e),
                        "fecha_escaneado": datetime.now().isoformat(),
                        "edicion": "",
                        "disciplina": "",
                        "recursos_requeridos": "",
                        "elementos_proteccion": "",
                        "descripcion_actividades": "",
                        "tipo_procedimiento": "",
                        "campo": ""
                    }
                    resultados.append(resultado_error)
        
        return resultados
    
    def get_generation_queue(self) -> List[Dict[str, Any]]:
        """
        Obtener cola actual de generación
        """
        resultado_scan = self.escanear_directorio()
        return resultado_scan.get("cola_generacion", [])
    
    def marcar_como_generado(self, codigo: str, version: str, preguntas_data: Dict[str, Any]):
        """
        Marcar un procedimiento como que ya tiene preguntas generadas
        """
        tracking_data = self.cargar_tracking_data()
        # Normalizar el código antes de construir la clave de tracking
        codigo_limpio = self._limpiar_codigo(codigo)
        tracking_key = create_tracking_key(codigo_limpio, version)
        
        # CORREGIDO: Manejar diferentes tipos de datos de preguntas
        preguntas_count = 0
        if isinstance(preguntas_data, dict):
            if "preguntas" in preguntas_data:
                preguntas_value = preguntas_data["preguntas"]
                if isinstance(preguntas_value, list):
                    preguntas_count = len(preguntas_value)
                elif isinstance(preguntas_value, int):
                    preguntas_count = preguntas_value
                else:
                    preguntas_count = 5  # Default
            elif "batch_id" in preguntas_data:
                preguntas_count = 5  # Default para lotes
            else:
                preguntas_count = 5  # Default
        elif isinstance(preguntas_data, list):
            preguntas_count = len(preguntas_data)
        elif isinstance(preguntas_data, int):
            preguntas_count = preguntas_data
        else:
            preguntas_count = 5  # Default
        
        tracking_data["generated_questions"][tracking_key] = {
            "codigo": codigo,
            "version": version,
            "fecha_generacion": datetime.now().isoformat(),
            "preguntas_count": preguntas_count,
            "status": "completed",
            "batch_info": preguntas_data if isinstance(preguntas_data, dict) else {"count": preguntas_count}
        }
        
        self.guardar_tracking_data(tracking_data)
        print(f"✅ Marcado como generado: {tracking_key} ({preguntas_count} preguntas)")
    
    def remover_de_cola(self, codigo: str, version: str) -> bool:
        """
        Remover un procedimiento de la cola de generación
        (marcándolo como procesado sin generar preguntas)
        """
        tracking_data = self.cargar_tracking_data()
        tracking_key = create_tracking_key(codigo, version)
        
        tracking_data["generated_questions"][tracking_key] = {
            "codigo": codigo,
            "version": version,
            "fecha_marcado": datetime.now().isoformat(),
            "status": "skipped",
            "reason": "removed_from_queue"
        }
        
        self.guardar_tracking_data(tracking_data)
        print(f"🚫 Removido de cola: {tracking_key}")
        return True

# =============================================================================
# FUNCIONES DE UTILIDAD
# =============================================================================

def crear_scanner(procedures_dir: str = None, max_workers: int = None) -> ProcedureScanner:
    """
    Crear instancia del scanner con configuración por defecto
    """
    if procedures_dir is None:
        procedures_dir = os.getenv(
            "PROCEDURES_SOURCE_DIR",
            str(get_admin_directory_path("procedures_source"))
        )
    
    if max_workers is None:
        # Determinar número óptimo de workers basado en CPU y archivos
        import multiprocessing
        max_workers = min(8, max(2, multiprocessing.cpu_count() // 2))
    
    return ProcedureScanner(procedures_dir, max_workers)

def escanear_procedimientos_rapido(procedures_dir: str = None) -> Dict[str, Any]:
    """
    Función de conveniencia para escaneo rápido
    """
    scanner = crear_scanner(procedures_dir)
    return scanner.escanear_directorio()

# =============================================================================
# TESTING
# =============================================================================

if __name__ == "__main__":
    # Test básico del scanner
    print("🧪 Testing ProcedureScanner...")
    
    scanner = crear_scanner()
    resultado = scanner.escanear_directorio()
    
    print(f"\n📊 Resultado del escaneo:")
    print(f"   - Success: {resultado['success']}")
    print(f"   - Archivos encontrados: {resultado['archivos_encontrados']}")
    print(f"   - Procedimientos nuevos: {resultado['procedimientos_nuevos']}")
    print(f"   - Procedimientos actualizados: {resultado['procedimientos_actualizados']}")
    print(f"   - Items en cola: {len(resultado['cola_generacion'])}")
    
    if resultado['cola_generacion']:
        print(f"\n📋 Primeros items en cola:")
        for i, item in enumerate(resultado['cola_generacion'][:3]):
            print(f"   {i+1}. {item['codigo']} v{item['version']} - {item['estado']}")