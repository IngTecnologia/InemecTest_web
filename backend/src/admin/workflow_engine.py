"""
Motor de workflow para orquestar el proceso completo de generación de preguntas
Integra: Scanner → Generator → Validators → Corrector → Excel Sync
"""

import os
import json
import asyncio
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Optional, Callable
from enum import Enum
import pandas as pd

from .models import (
    QuestionBatch,
    QueueItem,
    ProcedureStatus,
    QuestionStatus,
    ProcessingProgress,
    get_current_timestamp
)
from .config import (
    ADMIN_FILES,
    ADMIN_DIRECTORIES,
    WORKFLOW_STATES,
    MAX_PROCESSING_TIME_MINUTES,
    DEBUG_CONFIG,
    GENERATION_CONFIG
)
from .procedure_scanner import ProcedureScanner, crear_scanner
from .question_generator import QuestionGenerator, create_generator
from .validators import ValidationEngine, create_validation_engine
from .corrector import QuestionCorrector, create_corrector

class WorkflowState(str, Enum):
    """Estados del workflow"""
    IDLE = "idle"
    SCANNING = "scanning"
    QUEUED = "queued"
    GENERATING = "generating"
    VALIDATING = "validating"
    CORRECTING = "correcting"
    SYNCING = "syncing"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"

class ProcessingTask:
    """Tarea de procesamiento individual"""
    
    def __init__(self, queue_item: QueueItem, batch_id: str):
        self.queue_item = queue_item
        self.batch_id = batch_id
        self.state = WorkflowState.QUEUED
        self.current_step = 0
        self.total_steps = 5  # Scan, Generate, Validate, Correct, Sync
        self.started_at = get_current_timestamp()
        self.updated_at = get_current_timestamp()
        self.error_message = None
        self.question_batch = None
        self.processing_time_seconds = 0
        
    def update_progress(self, step: int, state: WorkflowState, message: str = ""):
        """Actualizar progreso de la tarea"""
        self.current_step = step
        self.state = state
        self.updated_at = get_current_timestamp()
        
        # Calcular tiempo de procesamiento
        start_time = datetime.fromisoformat(self.started_at)
        current_time = datetime.now()
        self.processing_time_seconds = (current_time - start_time).total_seconds()
        
        if DEBUG_CONFIG["verbose_logging"]:
            progress = (step / self.total_steps) * 100
            print(f"   📊 {self.queue_item.codigo}: {progress:.1f}% - {state.value} - {message}")
    
    def mark_failed(self, error: str):
        """Marcar tarea como fallida"""
        self.state = WorkflowState.FAILED
        self.error_message = error
        self.updated_at = get_current_timestamp()
    
    def mark_completed(self):
        """Marcar tarea como completada"""
        self.state = WorkflowState.COMPLETED
        self.current_step = self.total_steps
        self.updated_at = get_current_timestamp()

class WorkflowEngine:
    """
    Motor principal del workflow que orquesta todo el proceso
    """
    
    def __init__(self):
        """
        Inicializar motor de workflow
        """
        self.state = WorkflowState.IDLE
        self.processing_tasks: Dict[str, ProcessingTask] = {}
        self.active_batch_id = None
        self.progress_callbacks: List[Callable] = []
        
        # Inicializar componentes
        self._initialize_components()
        
        # Crear directorios necesarios
        self._ensure_directories()
        
        print(f"🔄 WorkflowEngine inicializado:")
        print(f"   - Estado inicial: {self.state}")
        print(f"   - Componentes cargados: Scanner, Generator, Validators, Corrector")
        print(f"   - Max tiempo procesamiento: {MAX_PROCESSING_TIME_MINUTES} min")
    
    def _initialize_components(self):
        """Inicializar todos los componentes del workflow"""
        try:
            self.scanner = crear_scanner()
            print("   ✅ Scanner inicializado")
        except Exception as e:
            print(f"   ❌ Error inicializando Scanner: {e}")
            self.scanner = None
        
        try:
            self.generator = create_generator()
            print("   ✅ Generator inicializado")
        except Exception as e:
            print(f"   ❌ Error inicializando Generator: {e}")
            print("   🔄 Activando modo mock para generator...")
            from .config import DEBUG_CONFIG
            DEBUG_CONFIG["mock_openai_calls"] = True
            # Intentar crear generador en modo mock
            try:
                self.generator = create_generator()
                print("   ✅ Generator inicializado en modo mock")
            except Exception as e2:
                print(f"   ❌ Error incluso en modo mock: {e2}")
                self.generator = None
        
        try:
            self.validation_engine = create_validation_engine()
            print("   ✅ ValidationEngine inicializado")
        except Exception as e:
            print(f"   ❌ Error inicializando ValidationEngine: {e}")
            self.validation_engine = None
        
        try:
            self.corrector = create_corrector()
            print("   ✅ Corrector inicializado")
        except Exception as e:
            print(f"   ❌ Error inicializando Corrector: {e}")
            self.corrector = None
    
    def _ensure_directories(self):
        """Crear directorios necesarios"""
        for dir_name, dir_path in ADMIN_DIRECTORIES.items():
            Path(dir_path).mkdir(parents=True, exist_ok=True)
    
    async def start_full_workflow(
        self, 
        procedure_codes: Optional[List[str]] = None,
        force_regeneration: bool = False
    ) -> str:
        """
        Iniciar workflow completo de generación
        
        Args:
            procedure_codes: Códigos específicos a procesar (None = todos en cola)
            force_regeneration: Forzar regeneración de preguntas existentes
            
        Returns:
            Batch ID del procesamiento iniciado
        """
        if self.state != WorkflowState.IDLE:
            raise ValueError(f"Workflow está ocupado. Estado actual: {self.state}")
        
        batch_id = f"workflow_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
        self.active_batch_id = batch_id
        
        print(f"🚀 Iniciando workflow completo - Batch ID: {batch_id}")
        
        try:
            # Paso 1: Escanear procedimientos
            self.state = WorkflowState.SCANNING
            await self._notify_progress("Escaneando procedimientos...")
            
            if not self.scanner:
                raise Exception("Scanner no disponible")
            
            scan_result = self.scanner.escanear_directorio()
            queue_items = scan_result["cola_generacion"]
            
            # Filtrar por códigos específicos si se proporcionan
            if procedure_codes:
                queue_items = [
                    item for item in queue_items 
                    if item["codigo"] in procedure_codes
                ]
            
            if not queue_items:
                print("   ⚠️ No hay procedimientos en cola para procesar")
                self.state = WorkflowState.IDLE
                return batch_id
            
            print(f"   📋 Procedimientos en cola: {len(queue_items)}")
            
            # Crear tareas de procesamiento
            self.processing_tasks = {}
            for item_data in queue_items:
                queue_item = QueueItem(**item_data)
                task = ProcessingTask(queue_item, batch_id)
                self.processing_tasks[queue_item.codigo] = task
            
            # Procesar cada procedimiento
            await self._process_all_procedures()
            
            self.state = WorkflowState.COMPLETED
            print(f"🎉 Workflow completado exitosamente - Batch ID: {batch_id}")
            
        except Exception as e:
            print(f"❌ Error en workflow: {e}")
            self.state = WorkflowState.FAILED
            await self._notify_progress(f"Error: {str(e)}")
            raise
        
        finally:
            self.active_batch_id = None
        
        return batch_id
    
    async def _process_all_procedures(self):
        """Procesar todos los procedimientos en las tareas"""
        total_procedures = len(self.processing_tasks)
        completed_procedures = 0
        failed_procedures = 0
        
        print(f"📊 Procesando {total_procedures} procedimientos...")
        
        for i, (codigo, task) in enumerate(self.processing_tasks.items()):
            try:
                print(f"\n📝 Procesando {i+1}/{total_procedures}: {codigo}")
                
                # Procesar procedimiento individual
                await self._process_single_procedure(task)
                
                if task.state == WorkflowState.COMPLETED:
                    completed_procedures += 1
                    print(f"   ✅ {codigo} completado exitosamente")
                else:
                    failed_procedures += 1
                    print(f"   ❌ {codigo} falló: {task.error_message}")
                
                # Actualizar progreso general
                overall_progress = f"Completados: {completed_procedures}, Fallidos: {failed_procedures}, Restantes: {total_procedures - i - 1}"
                await self._notify_progress(overall_progress)
                
                # Rate limiting entre procedimientos
                if i < total_procedures - 1:
                    await asyncio.sleep(2)
                
            except Exception as e:
                print(f"   ❌ Error crítico procesando {codigo}: {e}")
                task.mark_failed(str(e))
                failed_procedures += 1
        
        print(f"\n📊 Resumen final:")
        print(f"   - Total procesados: {total_procedures}")
        print(f"   - Exitosos: {completed_procedures}")
        print(f"   - Fallidos: {failed_procedures}")
        print(f"   - Tasa de éxito: {(completed_procedures/total_procedures)*100:.1f}%")
    
    async def _process_single_procedure(self, task: ProcessingTask):
        """Procesar un procedimiento individual a través de todo el pipeline"""
        codigo = task.queue_item.codigo
        
        try:
            # Paso 1: Preparar datos
            task.update_progress(1, WorkflowState.GENERATING, "Preparando generación...")
            
            procedure_file = Path(task.queue_item.datos_completos.ruta_completa)
            if not procedure_file.exists():
                raise FileNotFoundError(f"Archivo no encontrado: {procedure_file}")
            
            # Paso 2: Generar preguntas
            task.update_progress(2, WorkflowState.GENERATING, "Generando preguntas...")
            
            if not self.generator:
                # Intentar crear generador con modo mock activado
                print("   🔄 Intentando crear generador con modo mock...")
                from .config import DEBUG_CONFIG
                DEBUG_CONFIG["mock_openai_calls"] = True
                try:
                    self.generator = create_generator()
                    print("   ✅ Generator creado en modo mock")
                except Exception as e:
                    raise Exception(f"Generator no disponible ni en modo mock: {e}")
            
            question_batch = await self.generator.generate_questions_for_procedure(
                procedure_file, 
                task.queue_item.datos_completos.dict()
            )
            task.question_batch = question_batch
            
            print(f"   ✅ Preguntas generadas: {len(question_batch.questions)}")
            
            # Leer texto completo del procedimiento para validadores y corrector
            procedure_text = await self._extract_procedure_text(procedure_file)
            print(f"   📄 Texto del procedimiento extraído: {len(procedure_text)} caracteres")
            
            # Paso 3: Validar preguntas con procedimiento completo
            task.update_progress(3, WorkflowState.VALIDATING, "Validando preguntas...")
            
            if not self.validation_engine:
                raise Exception("ValidationEngine no disponible")
            
            validated_batch = await self.validation_engine.validate_batch(question_batch, procedure_text)
            task.question_batch = validated_batch
            
            validation_score = validated_batch.validation_score
            print(f"   ✅ Validación completada - Score: {validation_score:.2f}")
            
            # Paso 4: Corregir si es necesario con procedimiento completo
            task.update_progress(4, WorkflowState.CORRECTING, "Aplicando correcciones...")
            
            if not self.corrector:
                raise Exception("Corrector no disponible")
            
            corrected_batch = await self.corrector.correct_batch(validated_batch, procedure_text)
            task.question_batch = corrected_batch
            
            print(f"   ✅ Corrección completada")
            
            # Paso 5: Sincronizar con Excel (NUEVO)
            task.update_progress(5, WorkflowState.SYNCING, "Sincronizando con Excel...")
            
            # Importar el sync manager
            from .excel_sync import create_excel_sync_manager
            
            sync_manager = create_excel_sync_manager()
            sync_result = await sync_manager.sync_batch_to_excel(corrected_batch)
            
            if sync_result["success"]:
                print(f"   ✅ Sincronización con Excel completada")
            else:
                print(f"   ⚠️ Error en sincronización: {sync_result.get('error', 'Unknown')}")
            
            # Guardar resultados temporales
            await self._save_batch_results(corrected_batch)
            
            # NUEVO: Guardar en archivos finales
            print(f"🔄 [DEBUG] Iniciando guardado de archivos finales...")
            await self._save_to_final_files(corrected_batch)
            print(f"✅ [DEBUG] Guardado de archivos finales completado")
            
            # Marcar como completado
            task.mark_completed()
            
            # Actualizar tracking con más detalles
            if self.scanner:
                tracking_data = {
                    "batch_id": corrected_batch.batch_id, 
                    "preguntas": len(corrected_batch.questions),
                    "validation_score": validation_score,
                    "sync_success": sync_result["success"],
                    "excel_file": sync_result.get("excel_file", ""),
                    "completed_at": task.updated_at
                }
                self.scanner.marcar_como_generado(
                    codigo, 
                    task.queue_item.version,
                    tracking_data
                )
                print(f"✅ Marcado como generado: {codigo}_v{task.queue_item.version} ({len(corrected_batch.questions)} preguntas)")
            
            # NUEVO: Actualizar tracking principal
            print(f"🔄 [DEBUG] Iniciando actualización de tracking principal...")
            await self._update_main_tracking(corrected_batch, validation_score)
            print(f"✅ [DEBUG] Actualización de tracking principal completada")
            
        except Exception as e:
            task.mark_failed(str(e))
            raise
    
    async def _save_batch_results(self, batch: QuestionBatch):
        """Guardar resultados de un lote (temporal hasta tener excel_sync)"""
        try:
            # Crear directorio de resultados si no existe
            results_dir = Path(ADMIN_DIRECTORIES["temp"])
            results_dir.mkdir(parents=True, exist_ok=True)
            
            results_file = results_dir / f"{batch.batch_id}_results.json"
            
            # Convertir batch a diccionario serializable
            batch_data = {
                "batch_id": batch.batch_id,
                "procedure_codigo": batch.procedure_codigo,
                "procedure_version": batch.procedure_version,
                "procedure_name": batch.procedure_name,
                "status": batch.status.value,
                "created_at": batch.created_at,
                "updated_at": batch.updated_at,
                "total_questions": batch.total_questions,
                "validation_score": getattr(batch, 'validation_score', 0.0),
                "questions": []
            }
            
            # Agregar preguntas
            for question in batch.questions:
                question_data = {
                    "id": question.id,
                    "pregunta": question.pregunta,
                    "opciones": question.opciones,
                    "status": question.status.value,
                    "validations": [
                        {
                            "validator_type": v.validator_type.value,
                            "score": v.score,
                            "comment": v.comment,
                            "timestamp": v.timestamp
                        }
                        for v in question.validations
                    ],
                    "historial_revision": question.historial_revision,
                    # Campos adicionales para compatibilidad con Excel
                    "version_preg": getattr(question, 'version_preg', 1),
                    "prompt": getattr(question, 'prompt', "1.1"),
                    "puntaje_ia": getattr(question, 'puntaje_ia', 0)
                }
                batch_data["questions"].append(question_data)
            
            # Guardar archivo
            with open(results_file, 'w', encoding='utf-8') as f:
                json.dump(batch_data, f, indent=2, ensure_ascii=False)
            
            print(f"   💾 Resultados guardados: {results_file}")
            
            # OPCIONAL: También guardar en formato compatible con Excel original
            await self._save_to_excel_format(batch, results_file)
            
        except Exception as e:
            print(f"   ⚠️ Error guardando resultados: {e}")

    async def _save_to_excel_format(self, batch: QuestionBatch, json_file: Path):
        """Guardar en formato compatible con el Excel original (futuro)"""
        try:
            # Por ahora solo loggear, en el futuro aquí iría la sincronización con Excel
            excel_compatible_data = []
            
            for question in batch.questions:
                # Convertir al formato del Excel original
                excel_question = {
                    "codigo_procedimiento": batch.procedure_codigo,
                    "version_proc": int(batch.procedure_version),
                    "version_preg": getattr(question, 'version_preg', 1),
                    "prompt": getattr(question, 'prompt', "1.1"),
                    "puntaje_ia": getattr(question, 'puntaje_ia', 0),
                    "puntaje_e1": 0,
                    "puntaje_e2": 0,
                    "puntaje_e3": 0,
                    "puntaje_e4": 0,
                    "comentario_e1": "",
                    "comentario_e2": "",
                    "comentario_e3": "",
                    "comentario_e4": "",
                    "pregunta": question.pregunta,
                    "opciones": question.opciones,
                    "historial_revision": question.historial_revision
                }
                excel_compatible_data.append(excel_question)
            
            # Guardar versión compatible
            excel_file = json_file.parent / f"{batch.batch_id}_excel_format.json"
            with open(excel_file, 'w', encoding='utf-8') as f:
                json.dump(excel_compatible_data, f, indent=2, ensure_ascii=False)
            
            print(f"   📊 Formato Excel guardado: {excel_file}")
            
        except Exception as e:
            print(f"   ⚠️ Error guardando formato Excel: {e}")

    # 2. En procedure_scanner.py, AGREGAR este método:

    def marcar_como_procesado(self, codigo: str, version: str):
        """Marcar procedimiento como procesado en el tracking"""
        try:
            tracking_data = self.cargar_tracking_data()
            tracking_key = f"{codigo}_v{version}"
            
            # Actualizar o crear entrada
            tracking_data["generated_questions"][tracking_key] = {
                "codigo": codigo,
                "version": version,
                "fecha_procesado": datetime.now().isoformat(),
                "status": "completed",
                "questions_generated": True
            }
            
            self.guardar_tracking_data(tracking_data)
            print(f"✅ Procedimiento {tracking_key} marcado como procesado")
            
        except Exception as e:
            print(f"❌ Error marcando como procesado: {e}")

    # 3. En procedure_scanner.py, MODIFICAR el método determinar_estado_procedimiento:

    def determinar_estado_procedimiento(self, codigo: str, version: str, existing_codes: set) -> str:
        """Determinar el estado de un procedimiento"""
        tracking_data = self.cargar_tracking_data()
        tracking_key = f"{codigo}_v{version}"
        
        # Verificar si ya fue procesado
        if tracking_key in tracking_data.get("generated_questions", {}):
            status = tracking_data["generated_questions"][tracking_key].get("status")
            if status == "completed":
                return "ya_procesado"
        
        # Si el código no existe en Excel, es nuevo
        if codigo not in existing_codes:
            return "nuevo_procedimiento"
        
        # Si existe pero es nueva versión
        return "nueva_version"
    
    async def _notify_progress(self, message: str):
        """Notificar progreso a callbacks registrados"""
        for callback in self.progress_callbacks:
            try:
                await callback(self.state, message)
            except Exception as e:
                print(f"⚠️ Error en callback de progreso: {e}")
    
    def add_progress_callback(self, callback: Callable):
        """Agregar callback de progreso"""
        self.progress_callbacks.append(callback)
    
    def get_processing_progress(self, batch_id: Optional[str] = None) -> ProcessingProgress:
        """Obtener progreso actual del procesamiento"""
        if not self.processing_tasks or (batch_id and batch_id != self.active_batch_id):
            return ProcessingProgress(
                batch_id=batch_id or "no_active",
                total_procedures=0,
                completed=0,
                failed=0,
                current_step="No hay procesamiento activo",
                estimated_completion=None,
                updated_at=get_current_timestamp()
            )
        
        # Calcular progreso
        total = len(self.processing_tasks)
        completed = sum(1 for task in self.processing_tasks.values() if task.state == WorkflowState.COMPLETED)
        failed = sum(1 for task in self.processing_tasks.values() if task.state == WorkflowState.FAILED)
        
        # Determinar paso actual
        current_step = f"Procesando... ({completed + failed}/{total})"
        if self.state == WorkflowState.COMPLETED:
            current_step = "Procesamiento completado"
        elif self.state == WorkflowState.FAILED:
            current_step = "Procesamiento falló"
        
        return ProcessingProgress(
            batch_id=self.active_batch_id or "unknown",
            total_procedures=total,
            completed=completed,
            failed=failed,
            current_step=current_step,
            estimated_completion=None,
            updated_at=get_current_timestamp()
        )
    
    def get_workflow_stats(self) -> Dict[str, Any]:
        """Obtener estadísticas del workflow"""
        if not self.processing_tasks:
            return {"total_tasks": 0, "status": "idle"}
        
        stats = {
            "batch_id": self.active_batch_id,
            "workflow_state": self.state.value,
            "total_tasks": len(self.processing_tasks),
            "tasks_by_state": {},
            "average_processing_time": 0,
            "estimated_completion": None
        }
        
        # Contar por estado
        for state in WorkflowState:
            count = sum(1 for task in self.processing_tasks.values() if task.state == state)
            if count > 0:
                stats["tasks_by_state"][state.value] = count
        
        # Calcular tiempo promedio
        completed_tasks = [task for task in self.processing_tasks.values() if task.state == WorkflowState.COMPLETED]
        if completed_tasks:
            avg_time = sum(task.processing_time_seconds for task in completed_tasks) / len(completed_tasks)
            stats["average_processing_time"] = f"{avg_time/60:.1f} min"
        
        return stats
    
    async def _extract_procedure_text(self, procedure_file: Path) -> str:
        """
        Extraer el texto completo del procedimiento desde el archivo .docx
        """
        try:
            from docx import Document
            
            doc = Document(procedure_file)
            
            # Extraer todo el texto del documento
            full_text = []
            
            # Agregar texto de los párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Agregar texto de las tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            # Unir todo el texto
            procedure_text = "\n".join(full_text)
            
            print(f"   📄 Texto extraído del procedimiento: {len(procedure_text)} caracteres")
            
            return procedure_text
            
        except Exception as e:
            print(f"   ⚠️ Error extrayendo texto del procedimiento: {e}")
            # Retornar texto vacío en caso de error
            return ""

    def cancel_workflow(self) -> bool:
        """Cancelar workflow actual"""
        if self.state == WorkflowState.IDLE:
            return False
        
        print(f"🛑 Cancelando workflow {self.active_batch_id}")
        self.state = WorkflowState.CANCELLED
        
        # Marcar tareas como canceladas
        for task in self.processing_tasks.values():
            if task.state not in [WorkflowState.COMPLETED, WorkflowState.FAILED]:
                task.state = WorkflowState.CANCELLED
                task.updated_at = get_current_timestamp()
        
        self.active_batch_id = None
        return True

    async def _save_to_final_files(self, batch: QuestionBatch):
        """
        Guardar el lote en los archivos finales: generated_questions.json
        """
        try:
            print(f"💾 [DEBUG] Iniciando guardado de lote {batch.batch_id} en archivos finales...")
            
            # Importar funciones de configuración
            from .config import get_admin_file_path, BASE_DIR
            
            # DEBUG: Mostrar información del entorno
            import os
            print(f"💾 [DEBUG] Entorno actual:")
            print(f"   - ENVIRONMENT: {os.getenv('ENVIRONMENT', 'No definido')}")
            print(f"   - BASE_DIR: {BASE_DIR}")
            print(f"   - Directorio actual: {os.getcwd()}")
            print(f"   - PYTHONPATH: {os.getenv('PYTHONPATH', 'No definido')}")
            
            # Ruta del archivo de preguntas generadas
            generated_questions_file = get_admin_file_path("generated_questions")
            print(f"💾 [DEBUG] Ruta de archivo destino: {generated_questions_file}")
            print(f"💾 [DEBUG] Archivo existe: {generated_questions_file.exists()}")
            print(f"💾 [DEBUG] Directorio padre existe: {generated_questions_file.parent.exists()}")
            
            # Verificar permisos del directorio
            parent_dir = generated_questions_file.parent
            if parent_dir.exists():
                print(f"💾 [DEBUG] Permisos del directorio {parent_dir}: {oct(parent_dir.stat().st_mode)[-3:]}")
            else:
                print(f"💾 [DEBUG] Creando directorio padre: {parent_dir}")
                parent_dir.mkdir(parents=True, exist_ok=True)
            
            # Cargar preguntas existentes
            existing_questions = []
            if generated_questions_file.exists():
                try:
                    print(f"💾 [DEBUG] Leyendo archivo existente...")
                    with open(generated_questions_file, 'r', encoding='utf-8') as f:
                        existing_questions = json.load(f)
                    print(f"💾 [DEBUG] Preguntas existentes cargadas: {len(existing_questions)}")
                except json.JSONDecodeError as e:
                    print(f"⚠️ [DEBUG] Archivo generated_questions.json corrupto: {e}")
                    existing_questions = []
                except Exception as e:
                    print(f"❌ [DEBUG] Error leyendo archivo: {e}")
                    existing_questions = []
            else:
                print(f"💾 [DEBUG] Archivo no existe, será creado")
            
            # Convertir preguntas del batch al formato final
            print(f"💾 [DEBUG] Convirtiendo {len(batch.questions)} preguntas al formato final...")
            new_questions = []
            for i, question in enumerate(batch.questions):
                print(f"💾 [DEBUG] Procesando pregunta {i+1}: {question.pregunta[:50]}...")
                # DEBUG: Mostrar puntajes antes de convertir
                puntajes_debug = {
                    "puntaje_e1": getattr(question, 'puntaje_e1', 'NOT_FOUND'),
                    "puntaje_e2": getattr(question, 'puntaje_e2', 'NOT_FOUND'),
                    "puntaje_e3": getattr(question, 'puntaje_e3', 'NOT_FOUND'),
                    "puntaje_e4": getattr(question, 'puntaje_e4', 'NOT_FOUND')
                }
                print(f"💾 [DEBUG] Puntajes originales: {puntajes_debug}")
                question_data = {
                    "codigo_procedimiento": batch.procedure_codigo,
                    "version_proc": int(batch.procedure_version),
                    "version_preg": getattr(question, 'version_preg', 1),
                    "prompt": getattr(question, 'prompt', "1.1"),
                    "puntaje_ia": getattr(question, 'puntaje_ia', 0),
                    # CORREGIDO: Usar los puntajes reales de validación
                    "puntaje_e1": getattr(question, 'puntaje_e1', 0),
                    "puntaje_e2": getattr(question, 'puntaje_e2', 0),
                    "puntaje_e3": getattr(question, 'puntaje_e3', 0),
                    "puntaje_e4": getattr(question, 'puntaje_e4', 0),
                    "comentario_e1": getattr(question, 'comentario_e1', ""),
                    "comentario_e2": getattr(question, 'comentario_e2', ""),
                    "comentario_e3": getattr(question, 'comentario_e3', ""),
                    "comentario_e4": getattr(question, 'comentario_e4', ""),
                    "pregunta": question.pregunta,
                    "opciones": question.opciones,
                    "historial_revision": question.historial_revision,
                    # Metadatos adicionales
                    "batch_id": batch.batch_id,
                    "question_id": question.id,
                    "status": question.status.value,
                    "created_at": question.created_at,
                    "updated_at": question.updated_at
                }
                # DEBUG: Mostrar puntajes finales
                puntajes_finales = {
                    "puntaje_e1": question_data["puntaje_e1"],
                    "puntaje_e2": question_data["puntaje_e2"],
                    "puntaje_e3": question_data["puntaje_e3"],
                    "puntaje_e4": question_data["puntaje_e4"]
                }
                print(f"💾 [DEBUG] Puntajes guardados: {puntajes_finales}")
                
                new_questions.append(question_data)
            
            print(f"💾 [DEBUG] Preguntas convertidas: {len(new_questions)}")
            
            # Combinar con preguntas existentes
            all_questions = existing_questions + new_questions
            print(f"💾 [DEBUG] Total preguntas a guardar: {len(all_questions)}")
            
            # Guardar archivo actualizado
            print(f"💾 [DEBUG] Guardando archivo en: {generated_questions_file}")
            try:
                with open(generated_questions_file, 'w', encoding='utf-8') as f:
                    json.dump(all_questions, f, indent=2, ensure_ascii=False)
                print(f"💾 [DEBUG] Archivo guardado exitosamente")
                
                # Verificar que se guardó correctamente
                if generated_questions_file.exists():
                    file_size = generated_questions_file.stat().st_size
                    print(f"💾 [DEBUG] Verificación: archivo existe, tamaño: {file_size} bytes")
                else:
                    print(f"❌ [DEBUG] ERROR: Archivo no existe después de guardado")
                    
            except Exception as write_error:
                print(f"❌ [DEBUG] Error escribiendo archivo: {write_error}")
                raise
            
            print(f"   ✅ Guardadas {len(new_questions)} preguntas en {generated_questions_file}")
            print(f"   📊 Total preguntas en archivo: {len(all_questions)}")
            
        except Exception as e:
            print(f"   ❌ Error guardando en archivos finales: {e}")
            raise
    
    async def _update_main_tracking(self, batch: QuestionBatch, validation_score: float):
        """
        Actualizar el archivo principal de tracking: question_generation_tracking.json
        """
        try:
            print(f"📋 [DEBUG] Iniciando actualización de tracking para {batch.batch_id}...")
            
            # Importar funciones de configuración
            from .config import get_admin_file_path, BASE_DIR
            
            # DEBUG: Mostrar información del entorno
            import os
            print(f"📋 [DEBUG] Información de tracking:")
            print(f"   - BASE_DIR: {BASE_DIR}")
            print(f"   - Directorio actual: {os.getcwd()}")
            
            # Ruta del archivo de tracking
            tracking_file = get_admin_file_path("tracking")
            print(f"📋 [DEBUG] Ruta de archivo tracking: {tracking_file}")
            print(f"📋 [DEBUG] Archivo tracking existe: {tracking_file.exists()}")
            print(f"📋 [DEBUG] Directorio padre existe: {tracking_file.parent.exists()}")
            
            # Verificar/crear directorio padre
            parent_dir = tracking_file.parent
            if not parent_dir.exists():
                print(f"📋 [DEBUG] Creando directorio padre: {parent_dir}")
                parent_dir.mkdir(parents=True, exist_ok=True)
            
            # Cargar tracking existente
            tracking_data = {}
            if tracking_file.exists():
                try:
                    print(f"📋 [DEBUG] Leyendo archivo tracking existente...")
                    with open(tracking_file, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if content:
                            tracking_data = json.loads(content)
                            print(f"📋 [DEBUG] Tracking cargado: {len(tracking_data)} entradas")
                        else:
                            print(f"📋 [DEBUG] Archivo tracking está vacío")
                            tracking_data = {}
                except json.JSONDecodeError as e:
                    print(f"⚠️ [DEBUG] Archivo tracking corrupto: {e}")
                    tracking_data = {}
                except Exception as e:
                    print(f"❌ [DEBUG] Error leyendo tracking: {e}")
                    tracking_data = {}
            else:
                print(f"📋 [DEBUG] Archivo tracking no existe, será creado")
            
            # Crear estructura del batch para tracking
            batch_tracking = {
                "batch_id": batch.batch_id,
                "procedure_codigo": batch.procedure_codigo,
                "procedure_version": batch.procedure_version,
                "procedure_name": batch.procedure_name,
                "status": batch.status.value,
                "created_at": batch.created_at,
                "updated_at": batch.updated_at,
                "total_questions": len(batch.questions),
                "validation_score": validation_score,
                "questions": []
            }
            
            # Agregar detalles de cada pregunta
            for question in batch.questions:
                question_detail = {
                    "id": question.id,
                    "pregunta": question.pregunta,
                    "opciones": question.opciones,
                    "status": question.status.value,
                    "validations": [
                        {
                            "validator_type": v.validator_type.value,
                            "score": v.score,
                            "comment": v.comment,
                            "timestamp": v.timestamp
                        }
                        for v in question.validations
                    ],
                    "historial_revision": question.historial_revision,
                    "version_preg": getattr(question, 'version_preg', 1),
                    "prompt": getattr(question, 'prompt', "1.1"),
                    "puntaje_ia": getattr(question, 'puntaje_ia', 0)
                }
                batch_tracking["questions"].append(question_detail)
            
            # Actualizar tracking data con el nuevo batch
            # Usar el batch_id como clave para evitar duplicados
            print(f"📋 [DEBUG] Agregando batch al tracking: {batch.batch_id}")
            tracking_data[batch.batch_id] = batch_tracking
            print(f"📋 [DEBUG] Total entradas en tracking: {len(tracking_data)}")
            
            # Guardar archivo de tracking actualizado
            print(f"📋 [DEBUG] Guardando archivo tracking en: {tracking_file}")
            try:
                with open(tracking_file, 'w', encoding='utf-8') as f:
                    json.dump(tracking_data, f, indent=2, ensure_ascii=False)
                print(f"📋 [DEBUG] Archivo tracking guardado exitosamente")
                
                # Verificar que se guardó correctamente
                if tracking_file.exists():
                    file_size = tracking_file.stat().st_size
                    print(f"📋 [DEBUG] Verificación: archivo existe, tamaño: {file_size} bytes")
                else:
                    print(f"❌ [DEBUG] ERROR: Archivo tracking no existe después de guardado")
                    
            except Exception as write_error:
                print(f"❌ [DEBUG] Error escribiendo archivo tracking: {write_error}")
                raise
            
            print(f"   ✅ Tracking actualizado: {batch.batch_id}")
            print(f"   📊 Procedimiento: {batch.procedure_codigo} v{batch.procedure_version}")
            print(f"   📊 Preguntas: {len(batch.questions)}, Score: {validation_score:.2f}")
            
        except Exception as e:
            print(f"   ❌ Error actualizando tracking principal: {e}")
            raise

# =============================================================================
# FUNCIONES DE UTILIDAD
# =============================================================================

def create_workflow_engine() -> WorkflowEngine:
    """Crear instancia del motor de workflow"""
    return WorkflowEngine()

async def run_full_workflow(
    procedure_codes: Optional[List[str]] = None,
    force_regeneration: bool = False
) -> str:
    """Función de conveniencia para ejecutar workflow completo"""
    engine = create_workflow_engine()
    return await engine.start_full_workflow(procedure_codes, force_regeneration)

def enable_debug_workflow():
    """Habilitar modo debug para workflow"""
    DEBUG_CONFIG["enabled"] = True
    DEBUG_CONFIG["verbose_logging"] = True
    print("🧪 Modo debug de workflow habilitado")

# =============================================================================
# TESTING
# =============================================================================

async def test_workflow():
    """Testing completo del workflow engine"""
    print("🧪 Testing WorkflowEngine...")
    
    # Habilitar modo debug
    enable_debug_workflow()
    DEBUG_CONFIG["mock_openai_calls"] = True
    
    try:
        # Crear motor de workflow
        engine = create_workflow_engine()
        
        # Verificar inicialización
        print(f"✅ WorkflowEngine inicializado - Estado: {engine.state}")
        
        # Agregar callback de progreso
        async def progress_callback(state, message):
            print(f"📢 Callback: {state.value} - {message}")
        
        engine.add_progress_callback(progress_callback)
        
        # Simular workflow (solo si hay archivos de prueba)
        procedures_dir = Path(ADMIN_DIRECTORIES["procedures_source"])
        if procedures_dir.exists() and list(procedures_dir.glob("*.docx")):
            print("📁 Archivos de prueba encontrados, ejecutando workflow...")
            
            batch_id = await engine.start_full_workflow()
            print(f"✅ Workflow completado - Batch ID: {batch_id}")
            
            # Obtener estadísticas
            stats = engine.get_workflow_stats()
            print(f"📊 Estadísticas finales: {stats}")
            
        else:
            print("⚠️ No hay archivos de prueba, creando mock...")
            
            # Crear archivo de prueba mínimo
            procedures_dir.mkdir(parents=True, exist_ok=True)
            test_file = procedures_dir / "MOCK-001.docx"
            
            try:
                from docx import Document
                doc = Document()
                doc.add_paragraph("Procedimiento de prueba para workflow")
                doc.add_paragraph("1. Objetivo: Probar el sistema")
                doc.add_paragraph("5. Descripción de actividades")
                doc.save(test_file)
                print(f"✅ Archivo de prueba creado: {test_file}")
                
                # Ejecutar workflow de prueba
                batch_id = await engine.start_full_workflow()
                print(f"✅ Workflow de prueba completado - Batch ID: {batch_id}")
                
            except Exception as test_error:
                print(f"❌ Error en test: {test_error}")
        
    except Exception as e:
        print(f"❌ Error en testing: {e}")

if __name__ == "__main__":
    import asyncio
    asyncio.run(test_workflow())