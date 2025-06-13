"""
Generador de preguntas para el módulo administrativo de InemecTest
Adaptación del script generator.py original al framework del sistema
"""

import os
import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from openai import OpenAI
from .config import get_openai_api_key

from .models import (
    QuestionInProcess, 
    QuestionBatch, 
    ProcedureStatus, 
    QuestionStatus,
    GenerationConfig,
    get_current_timestamp
)
from .config import (
    get_system_message, 
    GENERATION_CONFIG, 
    DEBUG_CONFIG, 
    MOCK_RESPONSES,
    RATE_LIMIT_CONFIG
)

class QuestionGenerator:
    """
    Generador de preguntas usando OpenAI GPT-4o
    Basado en el script generator.py original con configuración centralizada
    """
    
    def __init__(self, config: Optional[GenerationConfig] = None):
        """
        Inicializar generador
        
        Args:
            config: Configuración de generación (usa defaults si no se proporciona)
        """
        self.config = config or GenerationConfig()
        
        # Obtener configuración centralizada
        self.generation_config = GENERATION_CONFIG
        self.debug_config = DEBUG_CONFIG
        self.rate_limit_config = RATE_LIMIT_CONFIG
        
        # CORREGIDO: Usar system_message en lugar de system_prompt
        self.system_message = get_system_message("generator")  # ← CORREGIDO
        
        # Inicializar cliente OpenAI solo si no estamos en modo mock
        if not self.debug_config["mock_openai_calls"]:
            api_key = get_openai_api_key()
            if not api_key:
                raise ValueError("OPENAI_API_KEY no está configurado en las variables de entorno")
            
            try:
                self.client = OpenAI(api_key=api_key)
                print("✅ Cliente OpenAI inicializado correctamente")
            except Exception as e:
                print(f"❌ Error inicializando cliente OpenAI: {e}")
                print("🔄 Habilitando modo mock automáticamente")
                self.debug_config["mock_openai_calls"] = True
                self.client = None
        else:
            self.client = None
            print("🧪 Modo DEBUG activado - usando respuestas mock")
        
        print(f"🤖 QuestionGenerator inicializado:")
        print(f"   - Modelo: {self.generation_config['openai_model']}")
        print(f"   - Max retries: {self.config.max_retries}")
        print(f"   - Timeout: {self.config.timeout_seconds}s")
        print(f"   - Debug mode: {self.debug_config['enabled']}")
        print(f"   - Mock calls: {self.debug_config['mock_openai_calls']}")
    
    def extraer_texto_docx(self, ruta_archivo: Path) -> str:
        """
        Extrae el texto completo de un archivo .docx
        Basado en tu script original
        """
        try:
            doc = Document(ruta_archivo)
            texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
            return texto
        except Exception as e:
            raise ValueError(f"Error extrayendo texto de {ruta_archivo}: {str(e)}")
    
    def _parse_codigo_version_from_filename(self, filename: str) -> Tuple[str, int]:
        """
        Extraer código y versión desde el nombre del archivo
        Compatible con lógica (comunicación/preguntas/generación/generator.py) original
        """
        # Remover extensión
        base_name = filename.replace('.docx', '').replace('.doc', '')
        
        # Detectar versión
        version = 1
        if "_V." in base_name:
            parts = base_name.split("_V.")
            if len(parts) == 2 and parts[1].isdigit():
                version = int(parts[1])
                codigo = parts[0]
            else:
                codigo = base_name
        else:
            codigo = base_name
        
        return codigo, version
    
    async def generate_questions_for_procedure(
        self, 
        procedure_file_path: Path, 
        procedure_data: Optional[Dict[str, Any]] = None
    ) -> QuestionBatch:
        """
        Generar 5 preguntas para un procedimiento específico
        
        Args:
            procedure_file_path: Ruta al archivo .docx del procedimiento
            procedure_data: Datos adicionales del procedimiento (opcional)
            
        Returns:
            QuestionBatch con las 5 preguntas generadas
        """
        try:
            print(f"🔄 Generando preguntas para: {procedure_file_path.name}")
            
            # Extraer contenido del archivo
            contenido = self.extraer_texto_docx(procedure_file_path)
            
            # Extraer código y versión del nombre del archivo
            codigo, version = self._parse_codigo_version_from_filename(procedure_file_path.name)
            
            # Crear mensaje para OpenAI
            mensaje_usuario = f"""
                Contenido del procedimiento:
                {contenido}

                Nombre del archivo: {procedure_file_path.name}
                """
            
            # Realizar llamada a OpenAI con reintentos
            preguntas_raw = await self._call_openai_with_retries(mensaje_usuario)
            
            # Parsear respuesta JSON
            preguntas_obj = json.loads(preguntas_raw)
            
            # Validar que se generaron exactamente 5 preguntas
            expected_questions = 1 if self.debug_config["test_with_single_question"] else 5
            if len(preguntas_obj) != expected_questions:
                if not self.debug_config["mock_openai_calls"]:
                    raise ValueError(f"Se esperaban {expected_questions} preguntas, se generaron {len(preguntas_obj)}")
                else:
                    # En modo mock, duplicar la pregunta para tener 5
                    while len(preguntas_obj) < 5:
                        pregunta_copy = preguntas_obj[0].copy()
                        pregunta_copy["pregunta"] = f"[MOCK {len(preguntas_obj)+1}] " + pregunta_copy["pregunta"]
                        preguntas_obj.append(pregunta_copy)
            
            # Crear objetos QuestionInProcess
            questions = []
            batch_id = f"batch_{codigo}_{version}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            for i, pregunta_data in enumerate(preguntas_obj):
                question_id = f"{batch_id}_q{i+1}"
                
                # Validar estructura de la pregunta
                self._validate_question_structure(pregunta_data)
                
                question = QuestionInProcess(
                    id=question_id,
                    procedure_codigo=codigo,
                    procedure_version=str(version),
                    pregunta=pregunta_data["pregunta"],
                    opciones=pregunta_data["opciones"],
                    version_preg=pregunta_data.get("version_preg", 1),
                    prompt=pregunta_data.get("prompt", "1.1"),
                    puntaje_ia=pregunta_data.get("puntaje_ia", 0),
                    validations=[],  # Se llenarán en el proceso de validación
                    status=QuestionStatus.generated,
                    created_at=get_current_timestamp(),
                    updated_at=get_current_timestamp(),
                    historial_revision=pregunta_data.get("historial_revision", [])
                )
                
                questions.append(question)
            
            # Crear QuestionBatch
            question_batch = QuestionBatch(
                batch_id=batch_id,
                procedure_codigo=codigo,
                procedure_version=str(version),
                procedure_name=procedure_data.get("nombre", "") if procedure_data else "",
                questions=questions,
                status=ProcedureStatus.generating,
                created_at=get_current_timestamp(),
                updated_at=get_current_timestamp(),
                total_questions=len(questions),
                questions_completed=len(questions),
                validation_score=0.0
            )
            
            print(f"✅ Preguntas generadas exitosamente para {codigo} v{version}")
            print(f"   - Batch ID: {batch_id}")
            print(f"   - Preguntas: {len(questions)}")
            
            return question_batch
            
        except json.JSONDecodeError as e:
            print(f"❌ Error parseando JSON de OpenAI: {e}")
            raise ValueError(f"Respuesta de OpenAI no es JSON válido: {str(e)}")
        except Exception as e:
            print(f"❌ Error generando preguntas: {e}")
            raise
    
    async def _call_openai_with_retries(self, mensaje_usuario: str) -> str:
        """
        Realizar llamada a OpenAI con manejo de reintentos y logging de respuesta
        """
        # Modo debug con mock responses
        if self.debug_config["mock_openai_calls"]:
            print("🧪 Usando respuesta mock para testing")
            import asyncio
            await asyncio.sleep(1)  # Simular latencia
            return MOCK_RESPONSES["generator"]
        
        if not self.client:
            raise ValueError("Cliente OpenAI no inicializado y no estamos en modo mock")
        
        last_error = None
        
        for attempt in range(self.config.max_retries):
            try:
                print(f"🤖 Llamada a OpenAI (intento {attempt + 1}/{self.config.max_retries})")
                
                respuesta = self.client.chat.completions.create(
                    model=self.generation_config["openai_model"],
                    messages=[
                        {
                            "role": "system",
                            "content": self.system_message  # ← CORREGIDO: era system_prompt
                        },
                        {
                            "role": "user",
                            "content": mensaje_usuario
                        }
                    ],
                    temperature=self.generation_config["temperature"],
                    max_tokens=self.generation_config["max_tokens"],
                    timeout=self.config.timeout_seconds
                )
                
                content = respuesta.choices[0].message.content
                if not content:
                    raise ValueError("OpenAI retornó contenido vacío")
                
                print(f"✅ Respuesta recibida de OpenAI ({len(content)} caracteres)")
                
                # ===== LOGGING DETALLADO DE LA RESPUESTA =====
                print("="*80)
                print("📋 RESPUESTA COMPLETA DE OPENAI:")
                print("="*80)
                print(content)
                print("="*80)
                print("🔍 ANÁLISIS DE LA RESPUESTA:")
                print(f"   📏 Longitud total: {len(content)} caracteres")
                print(f"   📐 Empieza con: '{content[:50]}...'")
                print(f"   📐 Termina con: '...{content[-50:]}'")
                
                # Verificar si es JSON válido
                try:
                    import json
                    parsed = json.loads(content.strip())
                    print(f"   ✅ JSON VÁLIDO - Tipo: {type(parsed)}")
                    if isinstance(parsed, list):
                        print(f"   📊 Lista con {len(parsed)} elementos")
                        if parsed:
                            print(f"   🔑 Claves del primer elemento: {list(parsed[0].keys())}")
                except json.JSONDecodeError as e:
                    print(f"   ❌ JSON INVÁLIDO: {e}")
                    print(f"   🔍 Error en posición: {getattr(e, 'pos', 'N/A')}")
                
                print("="*80)
                
                return content
                
            except Exception as e:
                last_error = e
                print(f"⚠️ Error en intento {attempt + 1}: {e}")
                
                if attempt < self.config.max_retries - 1:
                    import asyncio
                    wait_time = self.rate_limit_config["retry_delay_base"] ** attempt
                    wait_time = min(wait_time, self.rate_limit_config["max_retry_delay"])
                    print(f"🕒 Esperando {wait_time}s antes del siguiente intento...")
                    await asyncio.sleep(wait_time)
        
        raise Exception(f"Todos los intentos de llamada a OpenAI fallaron. Último error: {last_error}")
    
    def _validate_question_structure(self, pregunta_data: Dict[str, Any]) -> None:
        """
        Validar que una pregunta tenga la estructura correcta
        """
        required_fields = ["pregunta", "opciones"]
        
        for field in required_fields:
            if field not in pregunta_data:
                raise ValueError(f"Campo requerido faltante: {field}")
        
        if not isinstance(pregunta_data["opciones"], list):
            raise ValueError("'opciones' debe ser una lista")
        
        if len(pregunta_data["opciones"]) != 4:
            raise ValueError(f"Se esperan 4 opciones, se encontraron {len(pregunta_data['opciones'])}")
        
        if not pregunta_data["pregunta"].strip():
            raise ValueError("El texto de la pregunta no puede estar vacío")
        
        for i, opcion in enumerate(pregunta_data["opciones"]):
            if not isinstance(opcion, str) or not opcion.strip():
                raise ValueError(f"Opción {i+1} no es válida o está vacía")
        
        # Validación adicional en modo debug
        if self.debug_config["verbose_logging"]:
            print(f"✅ Pregunta validada: {pregunta_data['pregunta'][:50]}...")
    
    async def generate_questions_batch(
        self, 
        procedures_data: List[Dict[str, Any]]
    ) -> List[QuestionBatch]:
        """
        Generar preguntas para múltiples procedimientos
        
        Args:
            procedures_data: Lista de diccionarios con datos de procedimientos
                            Cada dict debe tener al menos 'ruta_completa' y opcionalmente otros datos
            
        Returns:
            Lista de QuestionBatch generados
        """
        batches = []
        total_procedures = len(procedures_data)
        
        print(f"🚀 Iniciando generación por lotes: {total_procedures} procedimientos")
        
        for i, proc_data in enumerate(procedures_data):
            try:
                print(f"📄 Procesando {i+1}/{total_procedures}: {proc_data.get('codigo', 'UNKNOWN')}")
                
                file_path = Path(proc_data["ruta_completa"])
                if not file_path.exists():
                    print(f"⚠️ Archivo no encontrado: {file_path}")
                    continue
                
                batch = await self.generate_questions_for_procedure(file_path, proc_data)
                batches.append(batch)
                
                print(f"✅ Completado {i+1}/{total_procedures}")
                
                # Rate limiting opcional
                if self.config.rate_limit_enabled and i < total_procedures - 1:
                    import asyncio
                    rate_limit_delay = 60 / self.rate_limit_config["requests_per_minute"]
                    print(f"🕒 Rate limiting: esperando {rate_limit_delay:.1f}s...")
                    await asyncio.sleep(rate_limit_delay)
                    
            except Exception as e:
                print(f"❌ Error procesando {proc_data.get('codigo', 'UNKNOWN')}: {e}")
                # Continuar con el siguiente procedimiento
                continue
        
        print(f"🎉 Generación por lotes completada: {len(batches)} lotes exitosos de {total_procedures} intentos")
        return batches
    
    def save_questions_to_json(
        self, 
        questions_batches: List[QuestionBatch], 
        output_file: Path
    ) -> None:
        """
        Guardar preguntas en formato JSON compatible con tu script original
        """
        try:
            all_questions = []
            
            for batch in questions_batches:
                for question in batch.questions:
                    # Convertir a formato compatible con tu script original
                    question_data = {
                        "codigo_procedimiento": question.procedure_codigo,
                        "version_proc": int(question.procedure_version),
                        "version_preg": question.version_preg,
                        "prompt": question.prompt,
                        "puntaje_ia": question.puntaje_ia,
                        "puntaje_e1": 0,
                        "puntaje_e2": 0,
                        "puntaje_e3": 0,
                        "puntaje_e4": 0,
                        "comentario_e1": "",
                        "comentario_e2": "",
                        "comentario_e3": "",
                        "comentario_e4": "",
                        "pregunta": question.pregunta,
                        "opciones": question.opciones,
                        "historial_revision": question.historial_revision
                    }
                    all_questions.append(question_data)
            
            # Guardar archivo JSON
            with open(output_file, "w", encoding="utf-8") as f:
                json.dump(all_questions, f, indent=2, ensure_ascii=False)
            
            print(f"💾 Preguntas guardadas en: {output_file}")
            print(f"📊 Total de preguntas: {len(all_questions)}")
            
        except Exception as e:
            print(f"❌ Error guardando preguntas: {e}")
            raise
    
    def get_generation_stats(self, batches: List[QuestionBatch]) -> Dict[str, Any]:
        """
        Obtener estadísticas de generación
        """
        if not batches:
            return {"total_batches": 0, "total_questions": 0, "success_rate": 0.0}
        
        total_questions = sum(len(batch.questions) for batch in batches)
        successful_batches = sum(1 for batch in batches if batch.status != ProcedureStatus.failed)
        
        stats = {
            "total_batches": len(batches),
            "successful_batches": successful_batches,
            "total_questions": total_questions,
            "success_rate": successful_batches / len(batches) if batches else 0.0,
            "average_questions_per_batch": total_questions / len(batches) if batches else 0.0,
            "procedures_processed": list(set(batch.procedure_codigo for batch in batches))
        }
        
        return stats

# =============================================================================
# FUNCIONES DE UTILIDAD
# =============================================================================

def create_generator(config: Optional[GenerationConfig] = None) -> QuestionGenerator:
    """
    Crear instancia del generador con configuración
    """
    return QuestionGenerator(config)

async def generate_questions_for_single_procedure(
    procedure_file: Path, 
    config: Optional[GenerationConfig] = None
) -> QuestionBatch:
    """
    Función de conveniencia para generar preguntas de un solo procedimiento
    """
    generator = create_generator(config)
    return await generator.generate_questions_for_procedure(procedure_file)

def enable_debug_mode():
    """Habilitar modo debug para testing"""
    DEBUG_CONFIG["enabled"] = True
    DEBUG_CONFIG["mock_openai_calls"] = True
    DEBUG_CONFIG["verbose_logging"] = True
    print("🧪 Modo debug habilitado")

def disable_debug_mode():
    """Deshabilitar modo debug"""
    DEBUG_CONFIG["enabled"] = False
    DEBUG_CONFIG["mock_openai_calls"] = False
    DEBUG_CONFIG["verbose_logging"] = False
    print("🔧 Modo debug deshabilitado")

# =============================================================================
# TESTING Y CLI
# =============================================================================

async def test_generation():
    """
    Función de testing para el generador
    """
    print("🧪 Testing QuestionGenerator...")
    
    # Habilitar modo debug para testing
    enable_debug_mode()
    
    # Verificar configuración
    from .config import validate_admin_config
    if not validate_admin_config():
        print("❌ Configuración inválida")
        return
    
    # Buscar archivos de prueba
    from .config import ADMIN_DIRECTORIES
    procedures_dir = Path(ADMIN_DIRECTORIES["procedures_source"])
    
    if not procedures_dir.exists():
        print(f"❌ Directorio no existe: {procedures_dir}")
        print("💡 Creando directorio y archivo de prueba...")
        procedures_dir.mkdir(parents=True, exist_ok=True)
        
        # Crear archivo de prueba
        test_file = procedures_dir / "TEST-001.docx"
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Este es un procedimiento de prueba.")
            doc.add_paragraph("Ítem 1: Objetivo del procedimiento")
            doc.add_paragraph("Ítem 5: Pasos de ejecución")
            doc.save(test_file)
            print(f"✅ Archivo de prueba creado: {test_file}")
        except Exception as e:
            print(f"❌ Error creando archivo de prueba: {e}")
            return
    
    docx_files = list(procedures_dir.glob("*.docx"))
    if not docx_files:
        print(f"❌ No se encontraron archivos .docx en {procedures_dir}")
        return
    
    # Tomar el primer archivo para prueba
    test_file = docx_files[0]
    print(f"📄 Archivo de prueba: {test_file}")
    
    try:
        # Crear generador
        generator = create_generator()
        
        # Generar preguntas
        batch = await generator.generate_questions_for_procedure(test_file)
        
        print(f"✅ Test exitoso!")
        print(f"   - Batch ID: {batch.batch_id}")
        print(f"   - Procedimiento: {batch.procedure_codigo} v{batch.procedure_version}")
        print(f"   - Preguntas generadas: {len(batch.questions)}")
        
        # Mostrar primera pregunta como ejemplo
        if batch.questions:
            first_q = batch.questions[0]
            print(f"   - Primera pregunta: {first_q.pregunta}")
            print(f"   - Opciones: {len(first_q.opciones)}")
            for i, opcion in enumerate(first_q.opciones):
                print(f"     {chr(65+i)}. {opcion}")
        
        # Guardar resultado para inspección
        output_file = Path("test_questions_output.json")
        generator.save_questions_to_json([batch], output_file)
        
        # Mostrar estadísticas
        stats = generator.get_generation_stats([batch])
        print(f"📊 Estadísticas: {stats}")
        
    except Exception as e:
        print(f"❌ Error en test: {e}")
        import traceback
        if DEBUG_CONFIG["verbose_logging"]:
            traceback.print_exc()

if __name__ == "__main__":
    import asyncio
    asyncio.run(test_generation())