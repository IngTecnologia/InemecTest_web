"""
API endpoints para el módulo administrativo
Integración completa con el workflow de generación de preguntas
"""

from fastapi import APIRouter, HTTPException, Query, BackgroundTasks, Depends, Header, UploadFile, File
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from datetime import datetime
import os
import asyncio

from .models import (
    AdminResponse, QueueResponse, GenerationStartResponse, ScanResult,
    QueueItem, ProcessingProgress, GenerationStats, BatchValidationSummary,
    BatchProcessingRequest, ProcedureStatus, ScannedProcedure, get_current_timestamp
)
from .procedure_scanner import ProcedureScanner, crear_scanner
from .workflow_engine import WorkflowEngine, create_workflow_engine
from .config import validate_admin_config, DEBUG_CONFIG
from .config import get_openai_api_key, validate_admin_credentials, get_admin_users

# Router para endpoints admin
admin_router = APIRouter(prefix="/admin", tags=["Admin - Generación de Preguntas"])

# Instancias globales
scanner_instance = None
workflow_engine_instance = None

# =============================================================================
# AUTENTICACIÓN
# =============================================================================

# Almacenar sesiones activas (en producción usar Redis o DB)
active_sessions = {}

class AdminLoginRequest(BaseModel):
    username: str
    code: str

class AdminLoginResponse(BaseModel):
    success: bool
    message: str
    user: Optional[Dict[str, Any]] = None
    session_token: Optional[str] = None

def verify_admin_session(authorization: str = Header(None)):
    """Middleware para verificar sesión de admin"""
    if not authorization:
        raise HTTPException(status_code=401, detail="No authorization header")
    
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization format")
    
    token = authorization.replace("Bearer ", "")
    
    if token not in active_sessions:
        raise HTTPException(status_code=401, detail="Invalid or expired session")
    
    return active_sessions[token]

@admin_router.post("/auth/login", response_model=AdminLoginResponse)
async def admin_login(request: AdminLoginRequest):
    """Login para acceso admin"""
    try:
        user_data = validate_admin_credentials(request.username, request.code)
        
        if not user_data:
            raise HTTPException(status_code=401, detail="Credenciales inválidas")
        
        # Guardar sesión
        session_token = user_data["session_token"]
        active_sessions[session_token] = user_data
        
        return AdminLoginResponse(
            success=True,
            message="Login exitoso",
            user={
                "username": user_data["username"],
                "name": user_data["name"],
                "permissions": user_data["permissions"]
            },
            session_token=session_token
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en login: {str(e)}")

@admin_router.post("/auth/logout")
async def admin_logout(current_user: Dict = Depends(verify_admin_session)):
    """Logout de admin"""
    # Buscar y eliminar sesión del usuario
    session_to_remove = None
    for token, user_data in active_sessions.items():
        if user_data["username"] == current_user["username"]:
            session_to_remove = token
            break
    
    if session_to_remove:
        del active_sessions[session_to_remove]
    
    return {"success": True, "message": "Logout exitoso"}

@admin_router.get("/auth/verify")
async def verify_session(current_user: Dict = Depends(verify_admin_session)):
    """Verificar si la sesión es válida"""
    return {
        "success": True,
        "user": {
            "username": current_user["username"],
            "name": current_user["name"],
            "permissions": current_user["permissions"]
        }
    }

@admin_router.get("/auth/users")
async def get_available_users():
    """Obtener lista de usuarios disponibles (solo usernames)"""
    users = get_admin_users()
    active_users = [username for username, user_data in users.items() if user_data.get("active", True)]
    return {"users": active_users}

# =============================================================================
# ENDPOINTS PROTEGIDOS (requieren autenticación)
# =============================================================================

@admin_router.post("/reset-instances")
async def reset_instances(current_user: Dict = Depends(verify_admin_session)):
    """Reiniciar instancias globales (para debug)"""
    global scanner_instance, workflow_engine_instance
    scanner_instance = None
    workflow_engine_instance = None
    return {"message": "Instancias reiniciadas"}

def get_scanner() -> ProcedureScanner:
    """Obtener instancia del scanner (singleton)"""
    global scanner_instance
    if scanner_instance is None:
        # Usar ruta absoluta para Docker
        if os.getenv("ENVIRONMENT") == "production":
            procedures_dir = "/app/data/procedures_source"
        else:
            procedures_dir = os.getenv("PROCEDURES_SOURCE_DIR", "data/procedures_source")
        
        print(f"🔍 [DEBUG] Inicializando scanner con directorio: {procedures_dir}")
        scanner_instance = crear_scanner(procedures_dir)
    return scanner_instance

def get_workflow_engine() -> WorkflowEngine:
    """Obtener instancia del workflow engine (singleton)"""
    global workflow_engine_instance
    if workflow_engine_instance is None:
        workflow_engine_instance = create_workflow_engine()
    return workflow_engine_instance

# =============================================================================
# ENDPOINTS DE ESTADO Y CONFIGURACIÓN
# =============================================================================

@admin_router.get("/status", response_model=AdminResponse)
async def get_admin_status(current_user: Dict = Depends(verify_admin_session)):
    """Estado general del módulo administrativo"""
    try:
        scanner = get_scanner()
        workflow_engine = get_workflow_engine()
        
        # Verificar estado de directorios y archivos
        procedures_dir_exists = scanner.procedures_source_dir.exists()
        tracking_file_exists = scanner.tracking_file.exists()
        excel_file_exists = scanner.excel_file.exists()
        
        # Cargar tracking data para estadísticas básicas
        tracking_data = scanner.cargar_tracking_data()
        
        # Estado del workflow
        workflow_stats = workflow_engine.get_workflow_stats()
        
        status_data = {
            "module": "admin",
            "version": "1.0.0",
            "workflow_state": workflow_engine.state.value,
            "active_batch_id": workflow_engine.active_batch_id,
            "procedures_source_dir": str(scanner.procedures_source_dir),
            "procedures_dir_exists": procedures_dir_exists,
            "tracking_file": str(scanner.tracking_file),
            "tracking_file_exists": tracking_file_exists,
            "excel_file": str(scanner.excel_file),
            "excel_file_exists": excel_file_exists,
            "generated_questions_count": len(tracking_data.get("generated_questions", {})),
            "last_scan": tracking_data.get("last_scan"),
            "workflow_stats": workflow_stats,
            "features": {
                "procedure_scanning": "active",
                "question_generation": "active",
                "validation_system": "active",
                "correction_engine": "active",
                "excel_sync": "pending",
                "workflow_engine": "active"
            },
            "environment": {
                "openai_api_key_set": bool(get_openai_api_key()),
                "admin_enabled": os.getenv("ADMIN_ENABLED", "false").lower() == "true",
                "debug_mode": DEBUG_CONFIG["enabled"],
                "mock_calls": DEBUG_CONFIG["mock_openai_calls"]
            }
        }
        
        return AdminResponse(
            success=True,
            message="Estado del módulo administrativo obtenido",
            data=status_data,
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo estado del módulo admin: {str(e)}"
        )

@admin_router.get("/config")
async def get_admin_config():
    """Obtener configuración actual del módulo admin"""
    try:
        from .config import (
            GENERATION_CONFIG, 
            VALIDATORS_CONFIG, 
            CORRECTOR_CONFIG,
            VALIDATION_THRESHOLD,
            get_enabled_validators
        )
        
        config_data = {
            "procedures_source_dir": os.getenv("PROCEDURES_SOURCE_DIR", "data/procedures_source"),
            "generation": GENERATION_CONFIG,
            "validators": {
                "enabled_validators": get_enabled_validators(),
                "validation_threshold": VALIDATION_THRESHOLD,
                "validators_config": VALIDATORS_CONFIG
            },
            "corrector": CORRECTOR_CONFIG,
            "debug": DEBUG_CONFIG,
            "environment": {
                "openai_api_key_set": bool(get_openai_api_key()),
                "generation_batch_size": int(os.getenv("GENERATION_BATCH_SIZE", "5")),
                "max_retries": int(os.getenv("MAX_RETRIES", "3"))
            }
        }
        
        return AdminResponse(
            success=True,
            message="Configuración del módulo admin obtenida",
            data=config_data,
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo configuración: {str(e)}"
        )

@admin_router.post("/config/debug")
async def toggle_debug_mode(enable: bool = Query(..., description="Habilitar/deshabilitar modo debug")):
    """Alternar modo debug del sistema - DESHABILITADO POR SEGURIDAD"""
    return AdminResponse(
        success=False,
        message="Modo debug deshabilitado por seguridad",
        data={"debug_enabled": False, "mock_calls": False},
        timestamp=get_current_timestamp()
    )

# =============================================================================
# ENDPOINTS DE ESCANEO Y COLA
# =============================================================================

@admin_router.post("/scan", response_model=ScanResult)
async def scan_procedures():
    """Escanear directorio de procedimientos y actualizar cola (OPTIMIZADO)"""
    try:
        import time
        start_time = time.time()
        
        scanner = get_scanner()
        resultado = scanner.escanear_directorio()
        
        scan_duration = time.time() - start_time
        print(f"🚀 [PERFORMANCE] Escaneo completado en {scan_duration:.2f} segundos")
        
        print(f"🔍 Debug - Resultado scanner: {resultado.keys()}")
        
        # Validar que el resultado tenga la estructura esperada
        if not isinstance(resultado, dict):
            raise ValueError("El scanner no devolvió un diccionario")
        
        required_keys = ["success", "message", "cola_generacion"]
        missing_keys = [key for key in required_keys if key not in resultado]
        if missing_keys:
            raise ValueError(f"El resultado del scanner no tiene las claves requeridas: {missing_keys}")
        
        # Asegurar que archivos_encontrados esté presente
        if "archivos_encontrados" not in resultado:
            resultado["archivos_encontrados"] = len(resultado.get("cola_generacion", []))
        
        # Convertir resultado a modelo Pydantic
        queue_items = []
        for i, item in enumerate(resultado["cola_generacion"]):
            print(f"🔍 Debug - Item {i}: {item.keys()}")
            print(f"🔍 Debug - datos_completos keys: {item['datos_completos'].keys()}")
            
            try:
                # Import local explícito para evitar problemas de contexto en threading
                from .models import ProcedureStatus as LocalProcedureStatus, ScannedProcedure as LocalScannedProcedure
                
                scanned_proc = LocalScannedProcedure(**item["datos_completos"])
                print(f"✅ ScannedProcedure creado para {item['codigo']}")
            except Exception as e:
                print(f"❌ Error creando ScannedProcedure: {e}")
                print(f"❌ Datos: {item['datos_completos']}")
                raise
            
            queue_item = QueueItem(
                codigo=item["codigo"],
                nombre=item["nombre"],
                version=item["version"],
                archivo=item["archivo"],
                estado=LocalProcedureStatus(item["estado"]),
                tracking_key=item["tracking_key"],
                datos_completos=scanned_proc,
                fecha_agregado=get_current_timestamp(),
                prioridad=1
            )
            queue_items.append(queue_item)
        
        return ScanResult(
            success=resultado["success"],
            message=f"{resultado['message']} (procesado en {scan_duration:.1f}s)",
            archivos_encontrados=resultado["archivos_encontrados"],
            procedimientos_nuevos=resultado.get("procedimientos_nuevos", 0),
            procedimientos_ya_procesados=resultado.get("procedimientos_ya_procesados", 0),
            total_procedimientos=resultado.get("total_procedimientos", len(queue_items)),
            cola_generacion=queue_items,
            tracking_file=str(scanner.tracking_file),
            timestamp=datetime.now().isoformat(),
            scan_duration=scan_duration
        )
        
    except Exception as e:
        print(f"❌ Error completo en scan: {e}")
        import traceback
        traceback.print_exc()
        
        # Devolver respuesta de error estructurada en lugar de HTTPException
        return ScanResult(
            success=False,
            message=f"Error escaneando procedimientos: {str(e)}",
            archivos_encontrados=0,
            procedimientos_nuevos=0,
            procedimientos_ya_procesados=0,
            total_procedimientos=0,
            cola_generacion=[],
            tracking_file="",
            timestamp=datetime.now().isoformat(),
            scan_duration=0.0
        )

@admin_router.get("/queue/simple")
async def get_generation_queue_simple():
    """Obtener lista simple de archivos en el directorio (para debug)"""
    try:
        import os
        if os.getenv("ENVIRONMENT") == "production":
            procedures_dir = "/app/data/procedures_source"
        else:
            procedures_dir = "data/procedures_source"
        
        from pathlib import Path
        dir_path = Path(procedures_dir)
        
        result = {
            "directory": str(dir_path),
            "exists": dir_path.exists(),
            "absolute_path": str(dir_path.absolute()),
            "files": [],
            "environment": os.getenv("ENVIRONMENT", "development"),
            "cwd": os.getcwd()
        }
        
        if dir_path.exists():
            files = list(dir_path.glob("*"))
            result["files"] = [{"name": f.name, "is_file": f.is_file(), "size": f.stat().st_size if f.is_file() else 0} for f in files]
        
        return result
        
    except Exception as e:
        return {"error": str(e), "traceback": str(e)}

@admin_router.get("/queue", response_model=QueueResponse)
async def get_generation_queue():
    """Obtener cola actual de generación"""
    try:
        print(f"🔍 [DEBUG] Iniciando obtención de cola...")
        
        scanner = get_scanner()
        print(f"🔍 [DEBUG] Scanner obtenido. Directorio: {scanner.procedures_source_dir}")
        print(f"🔍 [DEBUG] Directorio existe: {scanner.procedures_source_dir.exists()}")
        
        if scanner.procedures_source_dir.exists():
            archivos = list(scanner.procedures_source_dir.glob("*.docx"))
            print(f"🔍 [DEBUG] Archivos .docx encontrados: {len(archivos)}")
            for archivo in archivos:
                print(f"🔍 [DEBUG]   - {archivo.name}")
        else:
            print(f"❌ [DEBUG] Directorio no existe: {scanner.procedures_source_dir}")
        
        cola_raw = scanner.get_generation_queue()
        print(f"🔍 [DEBUG] Cola raw obtenida: {len(cola_raw)} items")
        
        # Convertir a modelos Pydantic
        queue_items = []
        status_summary = {}
        
        for i, item in enumerate(cola_raw):
            try:
                print(f"🔍 [DEBUG] Procesando item {i+1}: {item.get('codigo', 'UNKNOWN')}")
                
                # Import local explícito para evitar problemas de contexto en threading
                from .models import ProcedureStatus as LocalProcedureStatus, ScannedProcedure as LocalScannedProcedure
                
                queue_item = QueueItem(
                    codigo=item["codigo"],
                    nombre=item["nombre"],
                    version=item["version"],
                    archivo=item["archivo"],
                    estado=LocalProcedureStatus(item["estado"]),
                    tracking_key=item["tracking_key"],
                    datos_completos=LocalScannedProcedure(**item["datos_completos"]),
                    fecha_agregado=get_current_timestamp(),
                    prioridad=1
                )
                queue_items.append(queue_item)
                
                # Contar por estado
                estado = item["estado"]
                status_summary[estado] = status_summary.get(estado, 0) + 1
                
            except Exception as item_error:
                print(f"❌ [DEBUG] Error procesando item {i+1}: {item_error}")
                print(f"❌ [DEBUG] Item data: {item}")
                # Continuar con el siguiente item en lugar de fallar completamente
                continue
        
        # Obtener fecha del último escaneo
        try:
            tracking_data = scanner.cargar_tracking_data()
            last_scan = None
            if tracking_data.get("last_scan"):
                last_scan = tracking_data["last_scan"].get("timestamp")
            print(f"🔍 [DEBUG] Último escaneo: {last_scan}")
        except Exception as tracking_error:
            print(f"⚠️ [DEBUG] Error cargando tracking: {tracking_error}")
            last_scan = None
        
        print(f"🔍 [DEBUG] Queue items procesados: {len(queue_items)}")
        
        return QueueResponse(
            queue_items=queue_items,
            total_pending=len(queue_items),
            status_summary=status_summary,
            last_scan=last_scan
        )
        
    except Exception as e:
        print(f"❌ [DEBUG] Error completo en get_generation_queue: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo cola de generación: {str(e)}"
        )

@admin_router.delete("/queue/{codigo}/{version}")
async def remove_from_queue(codigo: str, version: str, reason: str = "manual_removal"):
    """Remover procedimiento específico de la cola"""
    try:
        scanner = get_scanner()
        success = scanner.remover_de_cola(codigo, version)
        
        if success:
            return AdminResponse(
                success=True,
                message=f"Procedimiento {codigo} v{version} removido de la cola",
                data={"codigo": codigo, "version": version, "reason": reason},
                timestamp=datetime.now().isoformat()
            )
        else:
            raise HTTPException(
                status_code=404,
                detail=f"Procedimiento {codigo} v{version} no encontrado en cola"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error removiendo procedimiento de cola: {str(e)}"
        )

# =============================================================================
# ENDPOINTS DE WORKFLOW Y GENERACIÓN
# =============================================================================

@admin_router.post("/workflow/start", response_model=GenerationStartResponse)
async def start_full_workflow(
    background_tasks: BackgroundTasks,
    request: Optional[BatchProcessingRequest] = None
):
    """Iniciar workflow completo de generación de preguntas"""
    try:
        workflow_engine = get_workflow_engine()
        
        # Verificar que el workflow no esté ocupado
        if workflow_engine.state.value != "idle":
            raise HTTPException(
                status_code=409,
                detail=f"Workflow está ocupado. Estado actual: {workflow_engine.state.value}"
            )
        
        # Determinar procedimientos a procesar
        procedure_codes = None
        if request and request.procedure_codes:
            procedure_codes = request.procedure_codes
        
        # Verificar que hay procedimientos para procesar
        scanner = get_scanner()
        cola = scanner.get_generation_queue()
        
        # Filtrar según si es regeneración o solo nuevos
        if procedure_codes:
            cola = [item for item in cola if item["codigo"] in procedure_codes]
        else:
            # Solo procesar procedimientos nuevos por defecto
            cola = [item for item in cola if item["estado"] in ["nuevo", "necesita_reproceso"]]
        
        if not cola:
            return GenerationStartResponse(
                batch_id="no_items",
                procedures_to_process=0,
                estimated_time_minutes=0,
                started_at=get_current_timestamp()
            )
        
        # Estimar tiempo (5 minutos por procedimiento como baseline)
        estimated_time = len(cola) * 5
        
        # Iniciar workflow en background
        async def run_workflow():
            try:
                batch_id = await workflow_engine.start_full_workflow(
                    procedure_codes=procedure_codes,
                    force_regeneration=request.force_regeneration if request else False
                )
                print(f"✅ Workflow completado - Batch ID: {batch_id}")
            except Exception as e:
                print(f"❌ Error en workflow background: {e}")
        
        background_tasks.add_task(run_workflow)
        
        # Generar batch_id temporal para tracking
        temp_batch_id = f"workflow_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        return GenerationStartResponse(
            batch_id=temp_batch_id,
            procedures_to_process=len(cola),
            estimated_time_minutes=estimated_time,
            started_at=get_current_timestamp()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error iniciando workflow: {str(e)}"
        )

@admin_router.post("/workflow/start-nuevos", response_model=GenerationStartResponse)
async def start_workflow_nuevos_only(
    background_tasks: BackgroundTasks
):
    """Iniciar workflow solo para procedimientos nuevos (Procesar Nuevos)"""
    try:
        workflow_engine = get_workflow_engine()
        
        # Verificar que el workflow no esté ocupado
        if workflow_engine.state.value != "idle":
            raise HTTPException(
                status_code=409,
                detail=f"Workflow está ocupado. Estado actual: {workflow_engine.state.value}"
            )
        
        # Obtener solo procedimientos nuevos
        scanner = get_scanner()
        cola = scanner.get_generation_queue()
        cola_nuevos = [item for item in cola if item["estado"] in ["nuevo", "necesita_reproceso"]]
        
        if not cola_nuevos:
            return GenerationStartResponse(
                batch_id="no_new_items",
                procedures_to_process=0,
                estimated_time_minutes=0,
                started_at=get_current_timestamp()
            )
        
        # Extraer códigos de procedimientos nuevos
        procedure_codes = [item["codigo"] for item in cola_nuevos]
        
        # Estimar tiempo (5 minutos por procedimiento como baseline)
        estimated_time = len(cola_nuevos) * 5
        
        # Iniciar workflow en background
        async def run_workflow():
            try:
                batch_id = await workflow_engine.start_full_workflow(
                    procedure_codes=procedure_codes,
                    force_regeneration=False
                )
                print(f"✅ Workflow para procedimientos nuevos completado - Batch ID: {batch_id}")
            except Exception as e:
                print(f"❌ Error en workflow background: {e}")
        
        background_tasks.add_task(run_workflow)
        
        # Generar batch_id temporal para tracking
        temp_batch_id = f"nuevos_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        return GenerationStartResponse(
            batch_id=temp_batch_id,
            procedures_to_process=len(cola_nuevos),
            estimated_time_minutes=estimated_time,
            started_at=get_current_timestamp()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error iniciando workflow para nuevos: {str(e)}"
        )

@admin_router.post("/workflow/regenerate/{codigo}/{version}", response_model=GenerationStartResponse)
async def regenerate_questions_for_procedure(
    codigo: str,
    version: str,
    background_tasks: BackgroundTasks,
    confirmed: bool = Query(False, description="Confirmación de regeneración")
):
    """Regenerar preguntas para un procedimiento específico (requiere confirmación)"""
    try:
        if not confirmed:
            raise HTTPException(
                status_code=400,
                detail="La regeneración requiere confirmación. Use el parámetro 'confirmed=true'"
            )
        
        workflow_engine = get_workflow_engine()
        
        # Verificar que el workflow no esté ocupado
        if workflow_engine.state.value != "idle":
            raise HTTPException(
                status_code=409,
                detail=f"Workflow está ocupado. Estado actual: {workflow_engine.state.value}"
            )
        
        # Verificar que el procedimiento existe y puede regenerarse
        scanner = get_scanner()
        cola = scanner.get_generation_queue()
        
        procedure_item = None
        for item in cola:
            if item["codigo"] == codigo and item["version"] == version:
                procedure_item = item
                break
        
        if not procedure_item:
            raise HTTPException(
                status_code=404,
                detail=f"Procedimiento {codigo} v{version} no encontrado"
            )
        
        if not procedure_item.get("puede_regenerar", False):
            raise HTTPException(
                status_code=400,
                detail=f"Procedimiento {codigo} v{version} no puede regenerarse. Estado: {procedure_item.get('estado')}"
            )
        
        # Iniciar workflow de regeneración en background
        async def run_regeneration():
            try:
                batch_id = await workflow_engine.start_full_workflow(
                    procedure_codes=[codigo],
                    force_regeneration=True
                )
                print(f"✅ Regeneración completada para {codigo} v{version} - Batch ID: {batch_id}")
            except Exception as e:
                print(f"❌ Error en regeneración background: {e}")
        
        background_tasks.add_task(run_regeneration)
        
        # Generar batch_id temporal para tracking
        temp_batch_id = f"regen_{codigo}_{version}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        return GenerationStartResponse(
            batch_id=temp_batch_id,
            procedures_to_process=1,
            estimated_time_minutes=5,
            started_at=get_current_timestamp()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error iniciando regeneración: {str(e)}"
        )

@admin_router.get("/workflow/progress/{batch_id}")
async def get_workflow_progress(batch_id: str = None):
    """Obtener progreso del workflow actual o específico"""
    try:
        workflow_engine = get_workflow_engine()
        progress = workflow_engine.get_processing_progress(batch_id)
        
        return AdminResponse(
            success=True,
            message=f"Progreso del workflow",
            data=progress.dict(),
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo progreso del workflow: {str(e)}"
        )

@admin_router.get("/workflow/status")
async def get_workflow_status():
    """Obtener estado actual del workflow"""
    try:
        workflow_engine = get_workflow_engine()
        stats = workflow_engine.get_workflow_stats()
        
        workflow_status = {
            "state": workflow_engine.state.value,
            "active_batch_id": workflow_engine.active_batch_id,
            "is_processing": workflow_engine.state.value != "idle",
            "stats": stats
        }
        
        return AdminResponse(
            success=True,
            message="Estado del workflow obtenido",
            data=workflow_status,
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo estado del workflow: {str(e)}"
        )

@admin_router.post("/workflow/cancel")
async def cancel_workflow():
    """Cancelar workflow actual"""
    try:
        workflow_engine = get_workflow_engine()
        success = workflow_engine.cancel_workflow()
        
        if success:
            return AdminResponse(
                success=True,
                message="Workflow cancelado exitosamente",
                data={"cancelled_batch_id": workflow_engine.active_batch_id},
                timestamp=datetime.now().isoformat()
            )
        else:
            return AdminResponse(
                success=False,
                message="No hay workflow activo para cancelar",
                data={"state": workflow_engine.state.value},
                timestamp=datetime.now().isoformat()
            )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error cancelando workflow: {str(e)}"
        )

# =============================================================================
# ENDPOINTS DE COMPONENTES INDIVIDUALES
# =============================================================================

@admin_router.post("/generate/single/{codigo}")
async def generate_questions_single_procedure(
    codigo: str,
    background_tasks: BackgroundTasks,
    version: str = Query("1", description="Versión del procedimiento")
):
    """Generar preguntas para un procedimiento específico"""
    try:
        from .question_generator import create_generator
        from pathlib import Path
        
        # Buscar archivo del procedimiento
        scanner = get_scanner()
        cola = scanner.get_generation_queue()
        
        procedure_item = None
        for item in cola:
            if item["codigo"] == codigo and item["version"] == version:
                procedure_item = item
                break
        
        if not procedure_item:
            raise HTTPException(
                status_code=404,
                detail=f"Procedimiento {codigo} v{version} no encontrado en cola"
            )
        
        async def generate_single():
            try:
                generator = create_generator()
                file_path = Path(procedure_item["datos_completos"]["ruta_completa"])
                
                batch = await generator.generate_questions_for_procedure(
                    file_path, 
                    procedure_item["datos_completos"]
                )
                
                print(f"✅ Preguntas generadas para {codigo}: {len(batch.questions)}")
                
            except Exception as e:
                print(f"❌ Error generando preguntas para {codigo}: {e}")
        
        background_tasks.add_task(generate_single)
        
        return AdminResponse(
            success=True,
            message=f"Generación iniciada para {codigo} v{version}",
            data={"codigo": codigo, "version": version},
            timestamp=datetime.now().isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error iniciando generación individual: {str(e)}"
        )

@admin_router.post("/validate/batch")
async def validate_questions_batch(batch_data: Dict[str, Any]):
    """Validar un lote de preguntas (endpoint de testing)"""
    try:
        from .validators import create_validation_engine
        from .models import QuestionBatch, QuestionInProcess
        
        # Convertir datos a modelos
        questions = []
        for q_data in batch_data.get("questions", []):
            question = QuestionInProcess(**q_data)
            questions.append(question)
        
        if not questions:
            raise HTTPException(status_code=400, detail="No hay preguntas para validar")
        
        batch = QuestionBatch(
            batch_id=batch_data.get("batch_id", f"test_{get_current_timestamp()}"),
            procedure_codigo=batch_data.get("procedure_codigo", "TEST"),
            procedure_version=batch_data.get("procedure_version", "1"),
            procedure_name=batch_data.get("procedure_name", "Test Procedure"),
            questions=questions,
            status=ProcedureStatus.validating,
            created_at=get_current_timestamp(),
            updated_at=get_current_timestamp()
        )
        
        # Validar
        validation_engine = create_validation_engine()
        validated_batch = await validation_engine.validate_batch(batch)
        
        # Obtener resumen
        summary = validation_engine.get_validation_summary(validated_batch)
        
        return AdminResponse(
            success=True,
            message="Validación completada",
            data=summary,
            timestamp=datetime.now().isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error en validación: {str(e)}"
        )

# =============================================================================
# ENDPOINTS DE ESTADÍSTICAS Y MONITOREO
# =============================================================================

@admin_router.get("/stats", response_model=GenerationStats)
async def get_admin_stats():
    """Obtener estadísticas completas del módulo admin"""
    try:
        scanner = get_scanner()
        workflow_engine = get_workflow_engine()
        
        # Obtener datos de tracking
        tracking_data = scanner.cargar_tracking_data()
        generated_questions = tracking_data.get("generated_questions", {})
        
        # Obtener cola actual
        cola = scanner.get_generation_queue()
        
        # Contar por estados
        procedures_in_queue = len(cola)
        procedures_completed = len([
            q for q in generated_questions.values() 
            if q.get("status") == "completed"
        ])
        
        # Obtener stats del workflow
        workflow_stats = workflow_engine.get_workflow_stats()
        
        # Calcular estadísticas
        total_procedures_scanned = len(generated_questions) + procedures_in_queue
        
        stats = GenerationStats(
            total_procedures_scanned=total_procedures_scanned,
            procedures_in_queue=procedures_in_queue,
            procedures_generating=workflow_stats.get("tasks_by_state", {}).get("generating", 0),
            procedures_validating=workflow_stats.get("tasks_by_state", {}).get("validating", 0),
            procedures_correcting=workflow_stats.get("tasks_by_state", {}).get("correcting", 0),
            procedures_completed=procedures_completed,
            procedures_failed=workflow_stats.get("tasks_by_state", {}).get("failed", 0),
            total_questions_generated=procedures_completed * 5,  # 5 preguntas por procedimiento
            total_questions_validated=procedures_completed * 5,  # Aproximación
            total_questions_corrected=procedures_completed * 5,  # Aproximación
            avg_generation_time_minutes=5.0,  # Baseline
            avg_validation_score=0.85,  # Estimación
            success_rate_percentage=95.0,  # Estimación
            last_scan_date=tracking_data.get("last_scan", {}).get("timestamp"),
            last_generation_date=None  # TODO: Obtener fecha real
        )
        
        return stats
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo estadísticas del admin: {str(e)}"
        )

@admin_router.get("/results")
async def get_generation_results():
    """Obtener resultados de generaciones recientes"""
    try:
        from pathlib import Path
        from .config import ADMIN_DIRECTORIES
        
        results_dir = Path(ADMIN_DIRECTORIES["temp"])
        if not results_dir.exists():
            return AdminResponse(
                success=True,
                message="No hay resultados disponibles",
                data={"results": []},
                timestamp=datetime.now().isoformat()
            )
        
        # Buscar archivos de resultados
        result_files = list(results_dir.glob("*_results.json"))
        results = []
        
        for file_path in sorted(result_files, key=lambda x: x.stat().st_mtime, reverse=True)[:10]:
            try:
                import json
                with open(file_path, 'r', encoding='utf-8') as f:
                    result_data = json.load(f)
                    
                    # Resumen del resultado
                    summary = {
                        "batch_id": result_data.get("batch_id"),
                        "procedure_codigo": result_data.get("procedure_codigo"),
                        "procedure_name": result_data.get("procedure_name"),
                        "status": result_data.get("status"),
                        "total_questions": result_data.get("total_questions", 0),
                        "validation_score": result_data.get("validation_score", 0),
                        "created_at": result_data.get("created_at"),
                        "file_path": str(file_path)
                    }
                    results.append(summary)
                    
            except Exception as e:
                print(f"Error leyendo resultado {file_path}: {e}")
                continue
        
        return AdminResponse(
            success=True,
            message=f"Resultados obtenidos: {len(results)}",
            data={"results": results, "total": len(results)},
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo resultados: {str(e)}"
        )

# =============================================================================
# ENDPOINTS DE GESTIÓN DE PRUEBAS PRESENTADAS
# =============================================================================

@admin_router.get("/evaluations/statistics")
async def get_evaluations_statistics(current_user: Dict = Depends(verify_admin_session)):
    """Obtener estadísticas de evaluaciones presentadas"""
    try:
        from ..excel_handler import ExcelHandler
        
        excel_handler = ExcelHandler()
        
        # Obtener evaluaciones con manejo de errores robusto
        try:
            evaluations = await excel_handler.get_all_evaluations()
            
            # Limpiar datos de evaluaciones para evitar objetos AdminResponse anidados
            if evaluations:
                clean_evaluations = []
                for eval_data in evaluations:
                    if isinstance(eval_data, dict):
                        # Convertir todos los valores a strings seguros
                        clean_eval = {}
                        for key, value in eval_data.items():
                            if isinstance(value, (str, int, float, bool)) or value is None:
                                clean_eval[key] = value
                            else:
                                # Convertir objetos complejos a string para evitar problemas
                                clean_eval[key] = str(value)
                        clean_evaluations.append(clean_eval)
                evaluations = clean_evaluations
                
        except Exception as excel_error:
            print(f"⚠️ Error leyendo Excel, devolviendo datos vacíos: {excel_error}")
            evaluations = []
        
        # Si no hay evaluaciones, retornar datos vacíos
        if not evaluations:
            return {
                "success": True,
                "message": "No hay evaluaciones registradas",
                "data": {
                    "total_evaluations": 0,
                    "by_campo": {},
                    "by_disciplina": {},
                    "approval_rate": 0,
                    "approved_count": 0,
                    "failed_count": 0,
                    "recent_evaluations": []
                },
                "timestamp": datetime.now().isoformat()
            }
        
        # Procesar estadísticas básicas
        stats_by_campo = {}
        approved_conocimiento_count = 0
        approved_aplicado_count = 0
        
        for eval_data in evaluations:
            # Procesar campo
            campo = eval_data.get("campo", "Sin campo")
            if campo:
                stats_by_campo[campo] = stats_by_campo.get(campo, 0) + 1
            
            # Contar aprobados conocimiento (automático ≥80%)
            aprobo_conocimiento = eval_data.get("aprobo_conocimiento", "")
            if str(aprobo_conocimiento).lower() in ["sí", "si", "yes", "true", "1"]:
                approved_conocimiento_count += 1
                
            # Contar aprobados conocimiento aplicado (manual supervisor)
            aprobo_aplicado = eval_data.get("aprobo", "")
            if str(aprobo_aplicado).lower() in ["sí", "si", "yes", "true", "1"]:
                approved_aplicado_count += 1
        
        conocimiento_rate = (approved_conocimiento_count / len(evaluations)) * 100 if evaluations else 0
        aplicado_rate = (approved_aplicado_count / len(evaluations)) * 100 if evaluations else 0
        
        # Preparar evaluaciones recientes de forma segura
        recent_evaluations = []
        try:
            sorted_evaluations = sorted(
                evaluations,
                key=lambda x: str(x.get("completed_at", "")),
                reverse=True
            )
            # Solo incluir campos básicos para evitar problemas de serialización
            # Limitar a 10 solo para evaluaciones recientes del dashboard
            for eval_data in sorted_evaluations[:10]:
                recent_evaluations.append({
                    "evaluation_id": str(eval_data.get("evaluation_id", "")),
                    "cedula": str(eval_data.get("cedula", "")),
                    "nombre": str(eval_data.get("nombre", "")),
                    "procedure_codigo": str(eval_data.get("procedure_codigo", "")),
                    "score_percentage": eval_data.get("score_percentage", 0),
                    "aprobo_conocimiento": str(eval_data.get("aprobo_conocimiento", "")),
                    "aprobo": str(eval_data.get("aprobo", "")),
                    "completed_at": str(eval_data.get("completed_at", ""))
                })
        except Exception as e:
            print(f"⚠️ Error procesando evaluaciones recientes: {e}")
            recent_evaluations = []
        
        return {
            "success": True,
            "message": "Estadísticas obtenidas exitosamente",
            "data": {
                "total_evaluations": len(evaluations),
                "by_campo": stats_by_campo,
                "by_disciplina": {"General": len(evaluations)},  # Simplificado
                "conocimiento": {
                    "approval_rate": round(conocimiento_rate, 2),
                    "approved_count": approved_conocimiento_count,
                    "failed_count": len(evaluations) - approved_conocimiento_count
                },
                "aplicado": {
                    "approval_rate": round(aplicado_rate, 2),
                    "approved_count": approved_aplicado_count,
                    "failed_count": len(evaluations) - approved_aplicado_count
                },
                "recent_evaluations": recent_evaluations
            },
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        print(f"❌ Error en get_evaluations_stats: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo estadísticas de evaluaciones: {str(e)}"
        )

@admin_router.get("/evaluations/stats")
async def get_evaluations_stats_alias(current_user: Dict = Depends(verify_admin_session)):
    """Alias para compatibilidad - estructura adaptada para frontend"""
    # Obtener datos completos
    full_response = await get_evaluations_statistics(current_user)
    
    # Adaptar estructura para compatibilidad con frontend actual
    if full_response.get("success") and "data" in full_response:
        data = full_response["data"]
        
        # Mantener estructura original pero agregar datos nuevos
        adapted_data = {
            "total_evaluations": data.get("total_evaluations", 0),
            "by_campo": data.get("by_campo", {}),
            "by_disciplina": data.get("by_disciplina", {}),
            # Usar aprobación de conocimiento como principal para compatibilidad
            "approval_rate": data.get("conocimiento", {}).get("approval_rate", 0),
            "approved_count": data.get("conocimiento", {}).get("approved_count", 0),
            "failed_count": data.get("conocimiento", {}).get("failed_count", 0),
            # Agregar datos nuevos
            "conocimiento": data.get("conocimiento", {}),
            "aplicado": data.get("aplicado", {}),
            "recent_evaluations": data.get("recent_evaluations", [])
        }
        
        return {
            "success": True,
            "message": full_response.get("message", ""),
            "data": adapted_data,
            "timestamp": full_response.get("timestamp", "")
        }
    
    return full_response

@admin_router.get("/evaluations/search")
async def search_evaluations(
    cedula: str = Query(None, description="Búsqueda por cédula"),
    campo: str = Query(None, description="Filtro por campo"),
    procedure_codigo: str = Query(None, description="Filtro por código de procedimiento"),
    limit: int = Query(None, description="Límite de resultados (None = todos)"),
    current_user: Dict = Depends(verify_admin_session)
):
    """Buscar evaluaciones con filtros"""
    try:
        from ..excel_handler import ExcelHandler
        
        excel_handler = ExcelHandler()
        evaluations = await excel_handler.get_all_evaluations()
        
        # Aplicar filtros
        filtered_evaluations = evaluations
        
        if cedula:
            filtered_evaluations = [
                e for e in filtered_evaluations 
                if cedula.lower() in e.get("cedula", "").lower()
            ]
        
        if campo:
            filtered_evaluations = [
                e for e in filtered_evaluations 
                if e.get("campo", "").lower() == campo.lower()
            ]
        
        if procedure_codigo:
            filtered_evaluations = [
                e for e in filtered_evaluations 
                if procedure_codigo.lower() in e.get("procedure_codigo", "").lower()
            ]
        
        # Ordenar por fecha más reciente
        filtered_evaluations = sorted(
            filtered_evaluations,
            key=lambda x: x.get("completed_at", ""),
            reverse=True
        )
        
        # Calcular estadísticas de los datos filtrados
        filtered_stats = {}
        if filtered_evaluations:
            approved_conocimiento_count = 0
            approved_aplicado_count = 0
            stats_by_campo = {}
            
            for eval_data in filtered_evaluations:
                # Procesar campo
                campo_eval = eval_data.get("campo", "Sin campo")
                if campo_eval:
                    stats_by_campo[campo_eval] = stats_by_campo.get(campo_eval, 0) + 1
                
                # Contar aprobados conocimiento (automático ≥80%)
                aprobo_conocimiento = eval_data.get("aprobo_conocimiento", "")
                if str(aprobo_conocimiento).lower() in ["sí", "si", "yes", "true", "1"]:
                    approved_conocimiento_count += 1
                    
                # Contar aprobados conocimiento aplicado (manual supervisor)
                aprobo_aplicado = eval_data.get("aprobo", "")
                if str(aprobo_aplicado).lower() in ["sí", "si", "yes", "true", "1"]:
                    approved_aplicado_count += 1
            
            conocimiento_rate = (approved_conocimiento_count / len(filtered_evaluations)) * 100
            aplicado_rate = (approved_aplicado_count / len(filtered_evaluations)) * 100
            
            filtered_stats = {
                "total_evaluations": len(filtered_evaluations),
                "by_campo": stats_by_campo,
                "conocimiento": {
                    "approval_rate": round(conocimiento_rate, 2),
                    "approved_count": approved_conocimiento_count,
                    "failed_count": len(filtered_evaluations) - approved_conocimiento_count
                },
                "aplicado": {
                    "approval_rate": round(aplicado_rate, 2),
                    "approved_count": approved_aplicado_count,
                    "failed_count": len(filtered_evaluations) - approved_aplicado_count
                }
            }
        
        # Aplicar límite para la respuesta solo si se especifica
        if limit is not None:
            result = filtered_evaluations[:limit]
        else:
            result = filtered_evaluations
        
        return AdminResponse.create_safe(
            success=True,
            message=f"Se encontraron {len(filtered_evaluations)} evaluaciones",
            data={
                "evaluations": result,
                "total_found": len(filtered_evaluations),
                "total_returned": len(result),
                "filtered_statistics": filtered_stats,
                "filters_applied": {
                    "cedula": cedula,
                    "campo": campo,
                    "procedure_codigo": procedure_codigo
                }
            }
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error buscando evaluaciones: {str(e)}"
        )

@admin_router.get("/evaluations/{evaluation_id}/report")
async def get_evaluation_report(
    evaluation_id: str,
    current_user: Dict = Depends(verify_admin_session)
):
    """Obtener reporte completo de una evaluación específica"""
    try:
        from ..excel_handler import ExcelHandler
        
        excel_handler = ExcelHandler()
        
        # Obtener datos principales de la evaluación
        evaluation_data = await excel_handler.get_evaluation_by_id(evaluation_id)
        if not evaluation_data:
            raise HTTPException(status_code=404, detail="Evaluación no encontrada")
        
        # Obtener respuestas detalladas
        answers = await excel_handler.get_evaluation_answers(evaluation_id)
        
        # Obtener conocimiento aplicado
        applied_knowledge = await excel_handler.get_evaluation_applied_knowledge(evaluation_id)
        
        # Obtener feedback
        feedback = await excel_handler.get_evaluation_feedback(evaluation_id)
        
        return AdminResponse(
            success=True,
            message="Reporte completo obtenido",
            data={
                "evaluation": evaluation_data,
                "answers": answers,
                "applied_knowledge": applied_knowledge,
                "feedback": feedback,
                "report_generated_at": get_current_timestamp()
            },
            timestamp=datetime.now().isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error obteniendo reporte de evaluación: {str(e)}"
        )

# =============================================================================
# ENDPOINTS DE UTILIDAD Y TESTING
# =============================================================================

@admin_router.post("/test/full-pipeline")
async def test_full_pipeline():
    """Endpoint de testing para todo el pipeline"""
    try:
        # Habilitar modo debug
        DEBUG_CONFIG["enabled"] = True
        DEBUG_CONFIG["mock_openai_calls"] = True
        DEBUG_CONFIG["verbose_logging"] = True
        
        workflow_engine = get_workflow_engine()
        
        # Verificar que el workflow esté libre
        if workflow_engine.state.value != "idle":
            raise HTTPException(
                status_code=409,
                detail="Workflow está ocupado, no se puede ejecutar test"
            )
        
        # Crear procedimiento de prueba si no existe
        from pathlib import Path
        from .config import ADMIN_DIRECTORIES
        
        procedures_dir = Path(ADMIN_DIRECTORIES["procedures_source"])
        procedures_dir.mkdir(parents=True, exist_ok=True)
        
        test_file = procedures_dir / "TEST-PIPELINE-001.docx"
        if not test_file.exists():
            try:
                from docx import Document
                doc = Document()
                doc.add_paragraph("PROCEDIMIENTO DE PRUEBA PARA PIPELINE COMPLETO")
                doc.add_paragraph("1. OBJETIVO: Probar el sistema completo de generación")
                doc.add_paragraph("5. PASOS DE EJECUCIÓN:")
                doc.add_paragraph("5.1 Verificar condiciones iniciales")
                doc.add_paragraph("5.2 Ejecutar procedimiento de prueba")
                doc.add_paragraph("5.3 Validar resultados")
                doc.save(test_file)
                print(f"✅ Archivo de prueba creado: {test_file}")
            except Exception as e:
                print(f"⚠️ Error creando archivo de prueba: {e}")
        
        # Ejecutar workflow de prueba
        batch_id = await workflow_engine.start_full_workflow(
            procedure_codes=["TEST-PIPELINE-001"]
        )
        
        return AdminResponse(
            success=True,
            message="Test completo del pipeline ejecutado",
            data={
                "batch_id": batch_id,
                "test_mode": True,
                "mock_calls": True,
                "procedure_tested": "TEST-PIPELINE-001"
            },
            timestamp=datetime.now().isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error en test del pipeline: {str(e)}"
        )

@admin_router.get("/health")
async def health_check_admin():
    """Health check específico del módulo admin"""
    try:
        # Verificar configuración
        config_valid = validate_admin_config()
        
        # Verificar componentes
        components_status = {}
        
        try:
            scanner = get_scanner()
            components_status["scanner"] = "ok"
        except Exception as e:
            components_status["scanner"] = f"error: {str(e)}"
        
        try:
            workflow_engine = get_workflow_engine()
            components_status["workflow_engine"] = workflow_engine.state.value
        except Exception as e:
            components_status["workflow_engine"] = f"error: {str(e)}"
        
        # Estado general
        overall_health = "healthy" if config_valid and all(
            "error" not in status for status in components_status.values()
        ) else "unhealthy"
        
        return AdminResponse(
            success=True,
            message=f"Health check completado - Estado: {overall_health}",
            data={
                "overall_health": overall_health,
                "config_valid": config_valid,
                "components": components_status,
                "openai_configured": bool(os.getenv("OPENAI_API_KEY")),
                "debug_mode": DEBUG_CONFIG["enabled"]
            },
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error en health check: {str(e)}"
        )

# =============================================================================
# ENDPOINTS DE CARGA DE PROCEDIMIENTOS
# =============================================================================

@admin_router.post("/procedures/upload")
async def upload_procedures(
    files: List[UploadFile] = File(...),
    current_user: Dict = Depends(verify_admin_session)
):
    """Cargar y validar procedimientos .docx con criterios específicos"""
    try:
        from pathlib import Path
        from .utils import extract_procedure_code_and_version
        import shutil
        
        # Validar que todos los archivos sean .docx
        for file in files:
            if not file.filename.endswith('.docx'):
                raise HTTPException(
                    status_code=400,
                    detail=f"Solo se aceptan archivos .docx. Archivo inválido: {file.filename}"
                )
        
        # Directorio de destino
        scanner = get_scanner()
        procedures_dir = scanner.procedures_source_dir
        procedures_dir.mkdir(parents=True, exist_ok=True)
        
        # Cargar tracking data actual para comparaciones
        tracking_data = scanner.cargar_tracking_data()
        generated_questions = tracking_data.get("generated_questions", {})
        
        # Obtener cola actual
        current_queue = scanner.get_generation_queue()
        
        results = []
        
        for file in files:
            result = {
                "filename": file.filename,
                "codigo": None,
                "version": None,
                "status": "error",
                "message": "Error desconocido",
                "details": None
            }
            
            try:
                # Criterio 1: Extraer código y versión del nombre del archivo
                try:
                    codigo, version = extract_procedure_code_and_version(file.filename)
                    result["codigo"] = codigo
                    result["version"] = version
                except Exception as e:
                    result["status"] = "error"
                    result["message"] = "Error extrayendo código y versión"
                    result["details"] = f"Formato de archivo inválido: {str(e)}"
                    results.append(result)
                    continue
                
                # Leer contenido del archivo para validaciones adicionales
                file_content = await file.read()
                await file.seek(0)  # Reset para uso posterior
                
                # Criterio 2: Validar contenido del archivo
                if len(file_content) == 0:
                    result["status"] = "error"
                    result["message"] = "Archivo vacío"
                    result["details"] = "El archivo no contiene datos"
                    results.append(result)
                    continue
                
                # Criterio 1 y 2: Validar código y versión internos del documento
                try:
                    # Guardar temporalmente el archivo para procesarlo
                    import tempfile
                    from docx import Document
                    
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                        temp_file.write(file_content)
                        temp_file_path = temp_file.name
                    
                    # Extraer datos del encabezado del documento
                    try:
                        doc = Document(temp_file_path)
                        header = doc.sections[0].header
                        tables = header.tables
                        
                        if not tables:
                            result["status"] = "error"
                            result["message"] = "Formato inválido"
                            result["details"] = "El documento no contiene tabla de encabezado"
                            results.append(result)
                            continue
                        
                        tabla = tables[0]
                        codigo_interno = tabla.cell(0, 2).text.upper().strip().replace("CÓDIGO:", "").replace("CODIGO:", "").strip()
                        version_interna = tabla.cell(1, 2).text.upper().strip().replace("VERSIÓN:", "").replace("VERSION:", "").strip()
                        
                        # Validar que el código del archivo coincida con el código interno
                        if codigo != codigo_interno:
                            result["status"] = "error"
                            result["message"] = "Código no coincide"
                            result["details"] = f"Código del archivo ({codigo}) no coincide con código interno ({codigo_interno})"
                            results.append(result)
                            continue
                        
                        # Validar que la versión del archivo coincida con la versión interna
                        try:
                            version_interna_int = int(version_interna)
                        except (ValueError, TypeError):
                            result["status"] = "error"
                            result["message"] = "Versión interna inválida"
                            result["details"] = f"La versión interna ({version_interna}) no es un número válido"
                            results.append(result)
                            continue
                        
                        if version != version_interna_int:
                            result["status"] = "error"
                            result["message"] = "Versión no coincide"
                            result["details"] = f"Versión del archivo ({version}) no coincide con versión interna ({version_interna_int})"
                            results.append(result)
                            continue
                        
                    finally:
                        # Limpiar archivo temporal
                        import os
                        try:
                            os.unlink(temp_file_path)
                        except:
                            pass
                            
                except Exception as e:
                    result["status"] = "error"
                    result["message"] = "Error validando documento"
                    result["details"] = f"Error procesando documento: {str(e)}"
                    results.append(result)
                    continue
                
                # Criterio 3: Verificar si ya existe un procedimiento con el mismo código
                existing_procedure = None
                existing_version = None
                
                # Buscar en preguntas generadas
                for tracking_key, question_data in generated_questions.items():
                    if question_data.get("codigo_procedimiento") == codigo:
                        existing_procedure = question_data
                        try:
                            existing_version = int(question_data.get("version_proc", 1))
                        except (ValueError, TypeError):
                            existing_version = 1
                        break
                
                # Buscar en cola actual
                if not existing_procedure:
                    for queue_item in current_queue:
                        if queue_item["codigo"] == codigo:
                            existing_procedure = queue_item
                            try:
                                existing_version = int(queue_item["version"])
                            except (ValueError, TypeError):
                                existing_version = 1
                            break
                
                # Criterio 4: Validar versión superior o código nuevo
                if existing_procedure:
                    if version <= existing_version:
                        result["status"] = "error"
                        result["message"] = "Versión no superior"
                        result["details"] = f"Ya existe versión {existing_version}. Nueva versión ({version}) debe ser superior."
                        results.append(result)
                        continue
                    else:
                        # Versión superior - aceptar
                        result["status"] = "success"
                        result["message"] = "Versión superior aceptada"
                        result["details"] = f"Actualización de versión {existing_version} → {version}"
                else:
                    # Código nuevo - aceptar
                    result["status"] = "success"
                    result["message"] = "Código nuevo aceptado"
                    result["details"] = "Procedimiento nuevo añadido al sistema"
                
                # Si llegamos aquí, el archivo es válido - copiarlo al directorio
                if result["status"] == "success":
                    dest_path = procedures_dir / file.filename
                    
                    # Escribir archivo
                    with open(dest_path, "wb") as buffer:
                        shutil.copyfileobj(file.file, buffer)
                    
                    result["details"] += f" | Archivo guardado en: {dest_path}"
                
            except Exception as e:
                result["status"] = "error"
                result["message"] = "Error procesando archivo"
                result["details"] = str(e)
            
            results.append(result)
        
        # Rescan automático después de cargar archivos exitosos
        successful_uploads = [r for r in results if r["status"] == "success"]
        if successful_uploads:
            try:
                # Ejecutar rescan en background
                async def rescan_after_upload():
                    await asyncio.sleep(1)  # Breve delay para asegurar que los archivos están listos
                    try:
                        scanner = get_scanner()
                        scan_result = scanner.escanear_directorio()
                        print(f"✅ Rescan automático completado: {scan_result.get('message', 'OK')}")
                    except Exception as e:
                        print(f"⚠️ Error en rescan automático: {e}")
                
                # Ejecutar rescan en background
                import threading
                thread = threading.Thread(target=lambda: asyncio.run(rescan_after_upload()))
                thread.daemon = True
                thread.start()
                
            except Exception as e:
                print(f"⚠️ Error iniciando rescan automático: {e}")
        
        return AdminResponse(
            success=True,
            message=f"Procesados {len(files)} archivos. {len(successful_uploads)} exitosos.",
            data={
                "results": results,
                "summary": {
                    "total_files": len(files),
                    "successful": len(successful_uploads),
                    "failed": len(results) - len(successful_uploads)
                }
            },
            timestamp=datetime.now().isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error procesando archivos: {str(e)}"
        )